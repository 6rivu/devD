"""
Standalone launcher for the interview practice feature only.
No PostgreSQL, no payments, no auth required.
"""

import os
import sys
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

os.environ["DB_HOST"] = "127.0.0.1"
os.environ["DB_PORT"] = "5432"
os.environ["DB_NAME"] = "cvolvepro"
os.environ["DB_USER"] = "postgres"
os.environ["DB_PASSWORD"] = ""
os.environ["AZURE_SPEECH_KEY"] = "ed57180d0ed84b19af73d6dcdcfc046c"
os.environ["AZURE_SPEECH_REGION"] = "eastus"

import streamlit as st
from io import BytesIO
from interview_module import show_interview_practice_page, DURATION_CREDITS, _init_session

st.set_page_config(
    page_title="CVOLVE PRO - AI Interview Practice",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

if "user_data" not in st.session_state:
    st.session_state.user_data = {"email": "demo@cvolvepro.com", "name": "Demo User", "credits": 100}
if "ai_model" not in st.session_state:
    st.session_state.ai_model = "gemini"
if "account_type" not in st.session_state:
    st.session_state.account_type = "individual"

_init_session()


def dummy_check_access(required_credits=2):
    return True


def dummy_deduct_credits(email, amount, feature=None):
    pass


def dummy_extract_resume(uploaded_file):
    if uploaded_file is None:
        return ""
    try:
        if uploaded_file.name.endswith(".pdf"):
            import PyPDF2 as pdf
            reader = pdf.PdfReader(uploaded_file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif uploaded_file.name.endswith(".docx"):
            from docx import Document
            doc = Document(uploaded_file)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        pass
    return "Extracted resume text (demo mode)"


def dummy_export_qa(qa_content):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt

    pdf_buf = BytesIO()
    doc = SimpleDocTemplate(pdf_buf)
    styles = getSampleStyleSheet()
    story = [Paragraph(qa_content.replace("\n", "<br/>"), styles["Normal"])]
    doc.build(story)
    pdf_buf.seek(0)

    docx_buf = BytesIO()
    wd = Document()
    for line in qa_content.split("\n"):
        if line.strip():
            wd.add_paragraph(line.strip())
    wd.save(docx_buf)
    docx_buf.seek(0)

    return pdf_buf, docx_buf


def dummy_generate_qa(resume_text, jd):
    return "Sample Q&A generated for: " + jd[:100] + "..."


st.markdown("""
<div style="background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:16px;padding:24px;margin-bottom:20px;">
    <h1 style="color:#e94560;margin:0;">🎯 CVOLVE PRO — Interview Practice</h1>
    <p style="color:#a0b4d6;margin:8px 0 0;">Standalone mode — no server setup needed.</p>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("### 💎 Demo Mode")
st.sidebar.info("⚡ Credits: 100 (unlimited in demo)")
st.sidebar.markdown("---")

show_interview_practice_page(
    check_access_fn=dummy_check_access,
    deduct_credits_fn=dummy_deduct_credits,
    extract_resume_fn=dummy_extract_resume,
    export_qa_fn=dummy_export_qa,
    generate_qa_fn=dummy_generate_qa,
)
