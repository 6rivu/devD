"""
CVOLVE PRO — AI Interview Practice Module
==========================================
Handles the full interview practice session:
 - Structured Q&A generation (Behavioral / Technical, 3 difficulty levels)
 - AI interviewer (question-by-question flow)
 - Voice answer recording + transcription (browser MediaRecorder via ST component)
 - AI evaluation (semantic, keyword, structure, completeness)
 - Downloadable feedback report (PDF + DOCX)
"""

import streamlit as st
import json
import re
import os
from io import BytesIO
from datetime import datetime

import google.generativeai as genai
import openai
from streamlit import session_state as st_session

# ─────────────────────────────────────────────────────────────────────────────
# Gemini model (shared with cv_generator)
# ─────────────────────────────────────────────────────────────────────────────
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
_model = genai.GenerativeModel("gemini-2.5-flash")
openai.api_key = os.getenv("OPENAI_API_KEY")


# ─────────────────────────────────────────────────────────────────────────────
# Credit costs per duration
# ─────────────────────────────────────────────────────────────────────────────
DURATION_CREDITS = {
    "15 minutes": 5,
    "30 minutes": 8,
    "45 minutes": 12,
}

# Questions per difficulty level per section (total Q per session depends on duration)
QUESTION_COUNTS = {
    "15 minutes": {"behavioral": {"Simple": 2, "Hard": 1, "Very Hard": 1}, "technical": {"Simple": 2, "Hard": 1, "Very Hard": 1}},
    "30 minutes": {"behavioral": {"Simple": 3, "Hard": 2, "Very Hard": 2}, "technical": {"Simple": 3, "Hard": 2, "Very Hard": 2}},
    "45 minutes": {"behavioral": {"Simple": 4, "Hard": 3, "Very Hard": 3}, "technical": {"Simple": 4, "Hard": 3, "Very Hard": 3}},
}


# ─────────────────────────────────────────────────────────────────────────────
# AI call helper (Gemini / OpenAI)
# ─────────────────────────────────────────────────────────────────────────────
def _ai_call(prompt: str, json_mode: bool = False) -> str:
    """Route to OpenAI or Gemini based on session setting."""
    if st_session.get("ai_model") == "openai":
        resp = openai.chat.completions.create(
            model="gpt-4.1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            response_format={"type": "json_object"} if json_mode else {"type": "text"},
        )
        return resp.choices[0].message.content
    else:
        resp = _model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.7,
                "response_mime_type": "application/json" if json_mode else "text/plain",
            },
        )
        return resp.text


# ─────────────────────────────────────────────────────────────────────────────
# 1. Generate structured interview Q&A
# ─────────────────────────────────────────────────────────────────────────────
def generate_structured_interview_qa(resume_text: str, job_description: str, duration: str) -> dict:
    """
    Returns a dict:
    {
        "behavioral": {
            "Simple":    [{"question": ..., "ideal_answer": ..., "key_points": [...]}, ...],
            "Hard":      [...],
            "Very Hard": [...]
        },
        "technical": { same structure }
    }
    """
    counts = QUESTION_COUNTS[duration]
    b = counts["behavioral"]
    t = counts["technical"]

    prompt = f"""
You are an expert technical interviewer and career coach.

Generate a structured interview question bank based on the candidate's resume and the job description below.

OUTPUT FORMAT (strict JSON only, no markdown, no extra text):
{{
  "behavioral": {{
    "Simple": [
      {{
        "question": "...",
        "ideal_answer": "A complete STAR-format answer (200-250 words)...",
        "key_points": ["key point 1", "key point 2", "key point 3", "key point 4", "key point 5"]
      }}
    ],
    "Hard": [...],
    "Very Hard": [...]
  }},
  "technical": {{
    "Simple": [...],
    "Hard": [...],
    "Very Hard": [...]
  }}
}}

RULES:
- Behavioral Simple: {b['Simple']} questions — standard STAR questions (leadership, teamwork, adaptability)
- Behavioral Hard: {b['Hard']} questions — complex situations with trade-offs and competing priorities
- Behavioral Very Hard: {b['Very Hard']} questions — executive-level leadership, ethics, ambiguity, influence without authority
- Technical Simple: {t['Simple']} questions — core concept checks, straightforward problem solving
- Technical Hard: {t['Hard']} questions — architecture decisions, system design, complex debugging, optimization
- Technical Very Hard: {t['Very Hard']} questions — open-ended design problems, scalability, cross-team trade-offs
- All questions must be grounded in the job description requirements and candidate's background
- ideal_answer should be a comprehensive, well-structured answer (200-250 words)
- key_points: 5-7 essential concepts/keywords the answer must cover to score well
- Do NOT include numbering inside question text

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        # Strip possible markdown fences
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        data = json.loads(raw_clean)
        return data
    except Exception as e:
        raise ValueError(f"Failed to parse AI response as JSON: {e}\n\nRaw: {raw[:500]}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. Flatten Q list for session (ordered: behavioral then technical)
# ─────────────────────────────────────────────────────────────────────────────
def flatten_questions(qa_bank: dict) -> list:
    """Returns flat list of question objects with section/difficulty metadata."""
    flat = []
    order = [
        ("behavioral", "Simple"),
        ("behavioral", "Hard"),
        ("behavioral", "Very Hard"),
        ("technical", "Simple"),
        ("technical", "Hard"),
        ("technical", "Very Hard"),
    ]
    for section, difficulty in order:
        qs = qa_bank.get(section, {}).get(difficulty, [])
        for q in qs:
            flat.append({
                "section": section,
                "difficulty": difficulty,
                "question": q.get("question", ""),
                "ideal_answer": q.get("ideal_answer", ""),
                "key_points": q.get("key_points", []),
            })
    return flat


# ─────────────────────────────────────────────────────────────────────────────
# 3. AI Answer Evaluation
# ─────────────────────────────────────────────────────────────────────────────
def evaluate_answer(question: str, ideal_answer: str, key_points: list,
                    user_answer: str, section: str, difficulty: str) -> dict:
    """
    Returns evaluation dict:
    {
      "score": 0-100,
      "meaning_match": 0-100,
      "keyword_coverage": 0-100,
      "keywords_covered": [...],
      "keywords_missed": [...],
      "structure_score": 0-100,
      "completeness_score": 0-100,
      "confidence_indicators": [...],
      "strengths": [...],
      "improvements": [...],
      "improved_answer": "...",
      "brief_feedback": "..."
    }
    """
    if not user_answer or len(user_answer.strip()) < 10:
        return {
            "score": 0, "meaning_match": 0, "keyword_coverage": 0,
            "keywords_covered": [], "keywords_missed": key_points,
            "structure_score": 0, "completeness_score": 0,
            "confidence_indicators": [],
            "strengths": [], "improvements": ["No answer was provided."],
            "improved_answer": ideal_answer,
            "brief_feedback": "No answer provided. Please attempt to answer all questions."
        }

    prompt = f"""
You are an expert interview coach evaluating a candidate's answer.

QUESTION: {question}
SECTION: {section} | DIFFICULTY: {difficulty}

IDEAL ANSWER (reference only — exact wording not required):
{ideal_answer}

KEY POINTS the answer should cover:
{json.dumps(key_points)}

CANDIDATE'S ANSWER:
{user_answer}

Evaluate the candidate's answer on these dimensions. Output ONLY valid JSON (no markdown):
{{
  "score": <overall 0-100>,
  "meaning_match": <0-100, how well the meaning aligns with the ideal>,
  "keyword_coverage": <0-100, % of key points addressed>,
  "keywords_covered": ["<key point>", ...],
  "keywords_missed": ["<missed key point>", ...],
  "structure_score": <0-100, logical flow, STAR or clear structure>,
  "completeness_score": <0-100, how complete and detailed the answer is>,
  "confidence_indicators": ["<positive indicator>", ...],
  "strengths": ["<strength 1>", "<strength 2>", ...],
  "improvements": ["<specific improvement>", ...],
  "improved_answer": "<a better version of the candidate's actual answer, 150-200 words>",
  "brief_feedback": "<2-3 sentence coach feedback>"
}}

SCORING GUIDELINES:
- Do NOT require exact wording match — evaluate semantic meaning
- A score of 80+ means the candidate covered the essential points well
- Confidence indicators: look for specific examples, numbers, confident language
- improved_answer should be based on the candidate's actual points, just enhanced
- Be fair but honest — do not inflate scores
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        return json.loads(raw_clean)
    except Exception:
        return {
            "score": 50, "meaning_match": 50, "keyword_coverage": 50,
            "keywords_covered": [], "keywords_missed": [],
            "structure_score": 50, "completeness_score": 50,
            "confidence_indicators": [],
            "strengths": ["Answer provided"], "improvements": ["Could not parse evaluation"],
            "improved_answer": ideal_answer,
            "brief_feedback": "Evaluation could not be parsed. Please try again."
        }


# ─────────────────────────────────────────────────────────────────────────────
# 4. Generate Full Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_feedback_report(session_results: list, duration: str) -> dict:
    """
    session_results: list of {question_obj, evaluation}
    Returns a structured feedback report dict.
    """
    total_score = 0
    behavioral_scores = []
    technical_scores = []
    all_keywords_covered = set()
    all_keywords_missed = set()
    well_answered = []
    incomplete_answers = []

    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        score = ev.get("score", 0)
        total_score += score

        if q_obj.get("section") == "behavioral":
            behavioral_scores.append(score)
        else:
            technical_scores.append(score)

        all_keywords_covered.update(ev.get("keywords_covered", []))
        all_keywords_missed.update(ev.get("keywords_missed", []))

        q_text = q_obj.get("question", "")
        if score >= 70:
            well_answered.append({"question": q_text, "score": score})
        else:
            incomplete_answers.append({
                "question": q_text,
                "score": score,
                "improvements": ev.get("improvements", []),
                "improved_answer": ev.get("improved_answer", ""),
            })

    n = len(session_results)
    overall = round(total_score / n) if n > 0 else 0
    behavioral_avg = round(sum(behavioral_scores) / len(behavioral_scores)) if behavioral_scores else 0
    technical_avg = round(sum(technical_scores) / len(technical_scores)) if technical_scores else 0

    # Remove keywords from missed if they were covered in other questions
    all_keywords_missed -= all_keywords_covered

    # Performance band
    if overall >= 85:
        band = "Excellent"
        band_color = "🟢"
    elif overall >= 70:
        band = "Good"
        band_color = "🔵"
    elif overall >= 55:
        band = "Average"
        band_color = "🟡"
    else:
        band = "Needs Improvement"
        band_color = "🔴"

    # Generate AI-powered insights
    results_summary = []
    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        results_summary.append({
            "question": q_obj.get("question", "")[:100],
            "section": q_obj.get("section", ""),
            "difficulty": q_obj.get("difficulty", ""),
            "score": ev.get("score", 0),
            "strengths": ev.get("strengths", []),
            "improvements": ev.get("improvements", []),
        })

    ai_prompt = f"""
You are a senior career coach reviewing a {duration} mock interview session.

OVERALL SCORE: {overall}/100 ({band})
BEHAVIORAL AVG: {behavioral_avg}/100
TECHNICAL AVG: {technical_avg}/100

SESSION RESULTS SUMMARY:
{json.dumps(results_summary, indent=2)}

Generate a concise, honest, and constructive performance report. Output ONLY valid JSON:
{{
  "overall_summary": "<2-3 sentence overall assessment>",
  "key_strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "weak_areas": ["<weak area 1>", "<weak area 2>"],
  "behavioral_feedback": "<specific feedback on behavioral answers>",
  "technical_feedback": "<specific feedback on technical answers>",
  "recommendations": ["<actionable recommendation 1>", "<recommendation 2>", "<recommendation 3>", "<recommendation 4>"],
  "next_steps": "<what to focus on in the next practice session>"
}}
"""

    try:
        raw = _ai_call(ai_prompt, json_mode=True)
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        ai_insights = json.loads(raw_clean)
    except Exception:
        ai_insights = {
            "overall_summary": f"You scored {overall}/100 overall. Review the individual question feedback to improve.",
            "key_strengths": ["Completed the practice session"],
            "weak_areas": ["Review questions scored below 70"],
            "behavioral_feedback": "Work on structuring answers using the STAR method.",
            "technical_feedback": "Strengthen technical depth with concrete examples.",
            "recommendations": [
                "Practice daily with different job descriptions",
                "Record yourself answering questions",
                "Research the company's tech stack thoroughly",
                "Prepare 3-4 strong STAR stories"
            ],
            "next_steps": "Focus on the questions where you scored below 70 and retry with improved answers."
        }

    return {
        "overall_score": overall,
        "performance_band": band,
        "band_color": band_color,
        "behavioral_score": behavioral_avg,
        "technical_score": technical_avg,
        "total_questions": n,
        "duration": duration,
        "well_answered": well_answered,
        "incomplete_answers": incomplete_answers,
        "keywords_covered": sorted(all_keywords_covered),
        "keywords_missed": sorted(all_keywords_missed),
        "session_results": session_results,
        **ai_insights,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 5. Export report to PDF & DOCX
# ─────────────────────────────────────────────────────────────────────────────
def export_feedback_report(report: dict):
    """Returns (pdf_buffer, docx_buffer)."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from docx import Document
    from docx.shared import Pt, RGBColor

    # ─── PDF ────────────────────────────────────────────────────────────────
    pdf_buf = BytesIO()
    doc = SimpleDocTemplate(pdf_buf, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("TitleS", fontSize=22, leading=28, fontName="Helvetica-Bold",
                              textColor=colors.HexColor("#1a1a2e"), spaceAfter=6)
    h2_s = ParagraphStyle("H2S", fontSize=14, leading=18, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#16213e"), spaceAfter=4, spaceBefore=12)
    h3_s = ParagraphStyle("H3S", fontSize=11, leading=14, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#0f3460"), spaceAfter=3, spaceBefore=8)
    body_s = ParagraphStyle("BodyS", fontSize=10, leading=14, spaceAfter=4)
    bullet_s = ParagraphStyle("BulletS", fontSize=10, leading=14, leftIndent=16, spaceAfter=3)
    score_s = ParagraphStyle("ScoreS", fontSize=28, fontName="Helvetica-Bold",
                             textColor=colors.HexColor("#e94560"), spaceAfter=4)
    sub_s = ParagraphStyle("SubS", fontSize=10, leading=13, textColor=colors.grey, spaceAfter=8)

    story = []

    # Header
    story.append(Paragraph("CVOLVE PRO — Interview Feedback Report", title_s))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}", sub_s))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#e94560")))

    # Score summary table
    story.append(Spacer(1, 10))
    score_data = [
        ["Overall Score", "Behavioral", "Technical", "Performance"],
        [f"{report['overall_score']}/100", f"{report['behavioral_score']}/100",
         f"{report['technical_score']}/100", report['performance_band']],
    ]
    score_table = Table(score_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
    score_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16213e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 11),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f4ff")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccddff")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(score_table)
    story.append(Spacer(1, 14))

    # Overall summary
    story.append(Paragraph("📋 Overall Assessment", h2_s))
    story.append(Paragraph(report.get("overall_summary", ""), body_s))

    # Strengths
    story.append(Paragraph("💪 Key Strengths", h2_s))
    for s in report.get("key_strengths", []):
        story.append(Paragraph(f"• {s}", bullet_s))

    # Weak areas
    story.append(Paragraph("⚠️ Areas to Improve", h2_s))
    for w in report.get("weak_areas", []):
        story.append(Paragraph(f"• {w}", bullet_s))

    # Behavioral & Technical feedback
    story.append(Paragraph("🎭 Behavioral Section Feedback", h2_s))
    story.append(Paragraph(report.get("behavioral_feedback", ""), body_s))
    story.append(Paragraph("⚙️ Technical Section Feedback", h2_s))
    story.append(Paragraph(report.get("technical_feedback", ""), body_s))

    # Questions answered well
    if report.get("well_answered"):
        story.append(Paragraph("✅ Questions Answered Well", h2_s))
        for item in report["well_answered"]:
            story.append(Paragraph(f"• {item['question']} — Score: {item['score']}/100", bullet_s))

    # Incomplete answers with improved versions
    if report.get("incomplete_answers"):
        story.append(Paragraph("📌 Questions Needing Improvement", h2_s))
        for item in report["incomplete_answers"]:
            story.append(Paragraph(item["question"], h3_s))
            story.append(Paragraph(f"Score: {item['score']}/100", sub_s))
            for imp in item.get("improvements", []):
                story.append(Paragraph(f"• {imp}", bullet_s))
            if item.get("improved_answer"):
                story.append(Paragraph("Suggested Improved Answer:", h3_s))
                story.append(Paragraph(item["improved_answer"], body_s))

    # Keywords
    story.append(Paragraph("🔑 Keywords Covered", h2_s))
    covered_text = ", ".join(report.get("keywords_covered", [])) or "None recorded"
    story.append(Paragraph(covered_text, body_s))

    story.append(Paragraph("❌ Keywords Missed", h2_s))
    missed_text = ", ".join(report.get("keywords_missed", [])) or "None — great coverage!"
    story.append(Paragraph(missed_text, body_s))

    # Recommendations
    story.append(Paragraph("🎯 Recommendations for Further Practice", h2_s))
    for r in report.get("recommendations", []):
        story.append(Paragraph(f"• {r}", bullet_s))

    story.append(Paragraph("🚀 Next Steps", h2_s))
    story.append(Paragraph(report.get("next_steps", ""), body_s))

    doc.build(story)
    pdf_buf.seek(0)

    # ─── DOCX ───────────────────────────────────────────────────────────────
    docx_buf = BytesIO()
    wd = Document()
    wd.core_properties.title = "CVOLVE PRO Interview Feedback Report"

    def add_h(doc, text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)

    def add_p(doc, text, bold=False, size=11):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    def add_b(doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    add_h(wd, "CVOLVE PRO — Interview Feedback Report", 1, (26, 26, 46))
    add_p(wd, f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}")
    add_h(wd, f"Overall Score: {report['overall_score']}/100 — {report['performance_band']}", 2, (233, 69, 96))
    add_p(wd, f"Behavioral: {report['behavioral_score']}/100 | Technical: {report['technical_score']}/100 | Questions: {report['total_questions']}")

    add_h(wd, "Overall Assessment", 2)
    add_p(wd, report.get("overall_summary", ""))

    add_h(wd, "Key Strengths", 2)
    for s in report.get("key_strengths", []):
        add_b(wd, s)

    add_h(wd, "Areas to Improve", 2)
    for w in report.get("weak_areas", []):
        add_b(wd, w)

    add_h(wd, "Behavioral Section Feedback", 2)
    add_p(wd, report.get("behavioral_feedback", ""))

    add_h(wd, "Technical Section Feedback", 2)
    add_p(wd, report.get("technical_feedback", ""))

    if report.get("well_answered"):
        add_h(wd, "Questions Answered Well", 2)
        for item in report["well_answered"]:
            add_b(wd, f"{item['question']} — {item['score']}/100")

    if report.get("incomplete_answers"):
        add_h(wd, "Questions Needing Improvement", 2)
        for item in report["incomplete_answers"]:
            add_h(wd, item["question"], 3)
            add_p(wd, f"Score: {item['score']}/100")
            for imp in item.get("improvements", []):
                add_b(wd, imp)
            if item.get("improved_answer"):
                add_p(wd, "Suggested Improved Answer:", bold=True)
                add_p(wd, item["improved_answer"])

    add_h(wd, "Keywords Covered", 2)
    add_p(wd, ", ".join(report.get("keywords_covered", [])) or "None")

    add_h(wd, "Keywords Missed", 2)
    add_p(wd, ", ".join(report.get("keywords_missed", [])) or "None — great coverage!")

    add_h(wd, "Recommendations", 2)
    for r in report.get("recommendations", []):
        add_b(wd, r)

    add_h(wd, "Next Steps", 2)
    add_p(wd, report.get("next_steps", ""))

    wd.save(docx_buf)
    docx_buf.seek(0)

    return pdf_buf, docx_buf


# ─────────────────────────────────────────────────────────────────────────────
# 6. Speech-to-text via browser component (JavaScript + Streamlit)
# ─────────────────────────────────────────────────────────────────────────────
VOICE_RECORDER_HTML = """
<style>
  .voice-btn {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 10px 20px; border-radius: 30px; border: none; cursor: pointer;
    font-size: 15px; font-weight: 600; transition: all 0.3s ease;
  }
  .voice-btn.record { background: linear-gradient(135deg, #e94560, #c0392b); color: #fff; }
  .voice-btn.stop   { background: linear-gradient(135deg, #2ecc71, #16a085); color: #fff; }
  .voice-btn:hover  { transform: scale(1.05); box-shadow: 0 4px 15px rgba(0,0,0,0.2); }
  #status { margin-top:10px; font-size:13px; color:#888; min-height:20px; }
  #transcript-box { 
    margin-top:12px; padding:12px; border-radius:10px; 
    background:#f8f9ff; border:1px solid #dee2ff; 
    font-size:14px; line-height:1.6; min-height:60px;
    white-space:pre-wrap;
  }
  .pulse { animation: pulse-anim 1s infinite; }
  @keyframes pulse-anim { 0%,100%{opacity:1} 50%{opacity:0.4} }
</style>

<div id="voice-widget">
  <button class="voice-btn record" id="startBtn" onclick="startRecording()">🎙️ Start Speaking</button>
  <button class="voice-btn stop" id="stopBtn" onclick="stopRecording()" style="display:none">⏹️ Stop Recording</button>
  <div id="status"></div>
  <div id="transcript-box" style="display:none"></div>
</div>

<script>
let mediaRecorder = null;
let audioChunks = [];
let stream = null;
let finalTranscript = '';
const API_URL = 'http://localhost:8000/api/transcribe';

function startRecording() {
  navigator.mediaDevices.getUserMedia({ audio: true })
    .then(s => {
      stream = s;
      audioChunks = [];
      const mime = MediaRecorder.isTypeSupported('audio/webm') ? 'audio/webm' : 'audio/ogg';
      mediaRecorder = new MediaRecorder(stream, { mimeType: mime });
      mediaRecorder.ondataavailable = (e) => { if (e.data.size > 0) audioChunks.push(e.data); };
      mediaRecorder.onstop = () => {
        const blob = new Blob(audioChunks, { type: mediaRecorder.mimeType });
        document.getElementById('status').textContent = '⏳ Transcribing...';
        fetch(API_URL, { method: 'POST', headers: { 'Content-Type': blob.type }, body: blob })
          .then(r => r.json())
          .then(data => {
            const txt = (data.transcript || '').trim();
            if (txt) {
              finalTranscript = txt;
              document.getElementById('status').textContent = '✅ Done!';
              window.parent.postMessage({type: 'voice_transcript', transcript: txt}, '*');
            } else {
              document.getElementById('status').textContent = data.error || 'No speech detected.';
            }
          })
          .catch(err => {
            document.getElementById('status').textContent = 'Error: ' + err.message;
          });
        if (stream) { stream.getTracks().forEach(t => t.stop()); stream = null; }
      };
      document.getElementById('status').innerHTML = '<span class="pulse">🔴 Recording... speak clearly</span>';
      document.getElementById('startBtn').style.display = 'none';
      document.getElementById('stopBtn').style.display = 'inline-flex';
      document.getElementById('transcript-box').style.display = 'block';
      mediaRecorder.start();
    })
    .catch(e => {
      document.getElementById('status').textContent = e.name === 'NotAllowedError' ? '❌ Mic blocked' : 'Error: ' + e.message;
    });
}

function stopRecording() {
  if (mediaRecorder && mediaRecorder.state !== 'inactive') mediaRecorder.stop();
  document.getElementById('startBtn').style.display = 'inline-flex';
  document.getElementById('stopBtn').style.display = 'none';
}
</script>
"""


# ─────────────────────────────────────────────────────────────────────────────
# 7. Main Streamlit UI — show_interview_practice_page()
# ─────────────────────────────────────────────────────────────────────────────
def show_interview_practice_page(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn):
    """
    Main entry point called from app.py.
    Passes down helper functions to avoid circular imports.
    """
    _init_session()

    # ── Phase router ──────────────────────────────────────────────────────────
    phase = st.session_state.get("interview_phase", "setup")

    if phase == "setup":
        _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn)
    elif phase == "session":
        _phase_session()
    elif phase == "report":
        _phase_report()


def _init_session():
    defaults = {
        "interview_phase": "setup",
        "interview_qa_bank": None,
        "interview_questions_flat": None,
        "interview_current_idx": 0,
        "interview_session_results": [],
        "interview_report": None,
        "interview_duration": "30 minutes",
        "interview_jd": "",
        "interview_resume_text": "",
        "voice_transcript_buffer": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────────────────────────────────────
# Phase 1: Setup
# ─────────────────────────────────────────────────────────────────────────────
def _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn):

    st.markdown("""
    <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
                border-radius: 16px; padding: 30px; margin-bottom: 24px;">
        <h1 style="color:#e94560; margin:0; font-size:32px;">🤖 AI Interview Practice</h1>
        <p style="color:#a0b4d6; margin:8px 0 0; font-size:16px;">
            A real-time interview simulation powered by AI — generate questions, practice answers,
            and receive a detailed performance report.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # ─── Credit Info Banner ───────────────────────────────────────────────────
    col_c1, col_c2, col_c3 = st.columns(3)
    with col_c1:
        st.markdown("""<div style="background:#fff3cd;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#856404">5 Credits</div>
            <div style="color:#856404;font-size:13px">15-minute session</div></div>""", unsafe_allow_html=True)
    with col_c2:
        st.markdown("""<div style="background:#d1ecf1;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#0c5460">8 Credits</div>
            <div style="color:#0c5460;font-size:13px">30-minute session</div></div>""", unsafe_allow_html=True)
    with col_c3:
        st.markdown("""<div style="background:#d4edda;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#155724">12 Credits</div>
            <div style="color:#155724;font-size:13px">45-minute session</div></div>""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ⚙️ Session Configuration")

    col1, col2 = st.columns([1, 1])

    with col1:
        duration = st.selectbox(
            "⏱️ Interview Duration",
            options=["15 minutes", "30 minutes", "45 minutes"],
            index=1,
            help="Longer sessions include more questions across all difficulty levels.",
            key="interview_duration_select"
        )
        st.session_state.interview_duration = duration

        jd = st.text_area(
            "📋 Job Description",
            height=200,
            placeholder="Paste the full job description here...",
            key="interview_jd_input",
            value=st.session_state.interview_jd,
        )
        st.session_state.interview_jd = jd

    with col2:
        uploaded = st.file_uploader(
            "📄 Upload Your Resume (PDF / DOCX)",
            type=["pdf", "docx"],
            key="interview_resume_upload"
        )

        st.markdown("#### 📊 What You'll Get")
        st.markdown("""
        <ul style="list-style:none;padding:0;margin:0">
          <li>✅ Behavioral questions (Simple → Very Hard)</li>
          <li>✅ Technical questions (Simple → Very Hard)</li>
          <li>✅ AI Interviewer — question by question</li>
          <li>✅ Type or speak your answers</li>
          <li>✅ AI evaluation on meaning, keywords, structure</li>
          <li>✅ Full feedback report (PDF + DOCX)</li>
          <li>✅ Suggested improved answers</li>
        </ul>
        """, unsafe_allow_html=True)

    # ─── Generate & Download Q&A Bank ────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 Step 1: Generate & Download Q&A Bank")
    st.markdown("*You can download the full Q&A bank (PDF/DOCX) before or after practicing.*")

    gen_col1, gen_col2 = st.columns([2, 1])

    if jd.strip() and uploaded:
        with gen_col1:
            if st.button("📚 Generate Q&A Bank + Start Practice Session",
                         type="primary", key="start_practice_btn",
                         use_container_width=True):
                credits_needed = DURATION_CREDITS[duration]
                if not check_access_fn(required_credits=credits_needed):
                    st.error(f"⚠️ You need {credits_needed} credits for a {duration} session. Please top up.")
                    return

                with st.spinner("🤖 AI is generating your personalized interview questions..."):
                    try:
                        resume_text = extract_resume_fn(uploaded)
                        st.session_state.interview_resume_text = resume_text

                        qa_bank = generate_structured_interview_qa(resume_text, jd, duration)
                        st.session_state.interview_qa_bank = qa_bank

                        flat = flatten_questions(qa_bank)
                        st.session_state.interview_questions_flat = flat
                        st.session_state.interview_current_idx = 0
                        st.session_state.interview_session_results = []
                        st.session_state.interview_report = None

                        # Deduct credits
                        user_email = st.session_state.user_data["email"]
                        deduct_credits_fn(user_email, credits_needed, feature="Interview Practice")

                        st.session_state.interview_phase = "session"
                        st.success(f"✅ {len(flat)} questions generated! {credits_needed} credits used. Starting session...")
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ Failed to generate questions: {str(e)}")

    else:
        st.info("👆 Please upload your resume and enter a job description to begin.")

    # Download-only option (if bank already exists)
    if st.session_state.interview_qa_bank:
        st.markdown("---")
        st.markdown("### 📁 Download Previously Generated Q&A Bank")
        bank = st.session_state.interview_qa_bank
        resume_text = st.session_state.interview_resume_text
        jd_text = st.session_state.interview_jd

        # Convert structured bank to flat text for export
        flat_text = _qa_bank_to_text(bank)
        try:
            pdf_buf, docx_buf = export_qa_fn(flat_text)
            dc1, dc2 = st.columns(2)
            with dc1:
                st.download_button("📥 Download Q&A PDF", data=pdf_buf,
                                   file_name="interview_qa_bank.pdf", mime="application/pdf",
                                   key="dl_qa_pdf_setup")
            with dc2:
                st.download_button("📥 Download Q&A DOCX", data=docx_buf,
                                   file_name="interview_qa_bank.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key="dl_qa_docx_setup")
        except Exception as e:
            st.warning(f"Export failed: {e}")


def _qa_bank_to_text(bank: dict) -> str:
    """Convert structured Q&A bank dict to displayable text for existing export_interview_qa."""
    lines = []
    for section in ["behavioral", "technical"]:
        section_title = "Behavioral Questions" if section == "behavioral" else "Technical Questions"
        lines.append(f"\n{section_title}\n")
        idx = 1
        for difficulty in ["Simple", "Hard", "Very Hard"]:
            qs = bank.get(section, {}).get(difficulty, [])
            if qs:
                lines.append(f"\n--- {difficulty} ---\n")
                for q in qs:
                    lines.append(f"{idx}. {q.get('question', '')}")
                    lines.append(f"Situation: {difficulty}-level scenario related to the role.")
                    lines.append(f"Task: Address the key requirements of the question.")
                    lines.append(f"Action: {q.get('ideal_answer', '')[:300]}...")
                    lines.append(f"Result: Demonstrate measurable impact and positive outcome.")
                    lines.append("")
                    idx += 1
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# Phase 2: Interview Session
# ─────────────────────────────────────────────────────────────────────────────
def _phase_session():
    questions = st.session_state.interview_questions_flat
    idx = st.session_state.interview_current_idx
    total = len(questions)

    if not questions:
        st.error("No questions found. Please go back and regenerate.")
        if st.button("↩️ Back to Setup"):
            st.session_state.interview_phase = "setup"
            st.rerun()
        return

    # Progress header
    progress_pct = idx / total
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:14px;padding:20px;margin-bottom:18px">
        <div style="display:flex;justify-content:space-between;align-items:center">
            <div>
                <h2 style="color:#e94560;margin:0">🤖 AI Interviewer</h2>
                <p style="color:#a0b4d6;margin:4px 0 0">Answer each question as if in a real interview</p>
            </div>
            <div style="text-align:right">
                <div style="color:#fff;font-size:20px;font-weight:700">Question {idx+1}/{total}</div>
                <div style="color:#a0b4d6;font-size:13px">Session in progress</div>
            </div>
        </div>
        <div style="margin-top:14px;background:#0f3460;border-radius:8px;height:8px;overflow:hidden">
            <div style="background:linear-gradient(90deg,#e94560,#fc5c7d);height:8px;width:{progress_pct*100:.1f}%;
                        border-radius:8px;transition:width 0.5s"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if idx >= total:
        # All questions answered
        _wrap_up_session()
        return

    q_obj = questions[idx]
    section = q_obj["section"].title()
    difficulty = q_obj["difficulty"]
    question_text = q_obj["question"]

    # Difficulty badge color
    diff_colors = {"Simple": "#2ecc71", "Hard": "#e67e22", "Very Hard": "#e94560"}
    diff_color = diff_colors.get(difficulty, "#888")

    # Question card
    st.markdown(f"""
    <div style="border-left:4px solid {diff_color};background:#f8f9ff;border-radius:0 12px 12px 0;
                padding:20px;margin-bottom:16px">
        <div style="display:flex;gap:10px;margin-bottom:10px;flex-wrap:wrap">
            <span style="background:#16213e;color:#a0b4d6;padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">
                {section}
            </span>
            <span style="background:{diff_color};color:#fff;padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">
                {difficulty}
            </span>
        </div>
        <p style="font-size:18px;font-weight:600;color:#1a1a2e;margin:0;line-height:1.5">{question_text}</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Answer input ──────────────────────────────────────────────────────────
    st.markdown("#### ✍️ Your Answer")

    tab_type, tab_voice = st.tabs(["⌨️ Type Answer", "🎙️ Speak Answer"])

    with tab_type:
        typed_answer = st.text_area(
            "Type your answer here:",
            height=200,
            placeholder="Give a structured answer. For behavioral questions, use STAR format (Situation, Task, Action, Result). For technical questions, explain your approach clearly.",
            key=f"typed_answer_{idx}",
            label_visibility="collapsed"
        )

    with tab_voice:
        st.markdown("""
        <div style="background:#fff3cd;border-radius:10px;padding:12px;margin-bottom:12px;font-size:13px;color:#856404">
            💡 <b>Tip:</b> Click "Start Speaking", answer the question aloud, then click "Stop Recording".
            The transcript will appear below. Then paste it above or submit directly.
        </div>
        """, unsafe_allow_html=True)

        st.components.v1.html(VOICE_RECORDER_HTML, height=200, scrolling=False)

        voice_answer = st.text_area(
            "Transcribed / edited voice answer:",
            height=120,
            placeholder="Your spoken answer will appear here after recording. You can also edit it.",
            key=f"voice_answer_{idx}",
            value=st.session_state.get("voice_transcript_buffer", ""),
            label_visibility="visible"
        )

    # Combine: prefer typed if both filled, else voice
    final_answer = typed_answer.strip() if typed_answer.strip() else voice_answer.strip()

    # ── Skip / Submit ─────────────────────────────────────────────────────────
    btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])

    with btn_col1:
        submit_disabled = not final_answer
        if st.button("✅ Submit Answer & Next Question",
                     type="primary", key=f"submit_{idx}",
                     disabled=submit_disabled,
                     use_container_width=True):
            _submit_answer(q_obj, final_answer)

    with btn_col2:
        if st.button("⏭️ Skip Question", key=f"skip_{idx}", use_container_width=True):
            _submit_answer(q_obj, "")  # empty = skipped

    with btn_col3:
        if st.button("🔚 End Session Early", key=f"end_{idx}", use_container_width=True):
            _wrap_up_session(early=True)

    # ── Hints panel ───────────────────────────────────────────────────────────
    with st.expander("💡 Hints — Key Points to Cover", expanded=False):
        st.markdown("*These are the key concepts your answer should address (shown for practice purposes):*")
        for kp in q_obj.get("key_points", []):
            st.markdown(f"• {kp}")

    # ── Previous results quick view ───────────────────────────────────────────
    if st.session_state.interview_session_results:
        with st.expander(f"📋 Previous Answers — {len(st.session_state.interview_session_results)} answered", expanded=False):
            for i, res in enumerate(st.session_state.interview_session_results):
                ev = res.get("evaluation", {})
                score = ev.get("score", 0)
                color = "#2ecc71" if score >= 70 else "#e67e22" if score >= 50 else "#e94560"
                q_text = res["question_obj"]["question"][:80]
                st.markdown(f"""
                <div style="display:flex;justify-content:space-between;padding:8px;
                            border-radius:8px;background:#f8f9ff;margin-bottom:6px">
                    <span style="font-size:13px;color:#333">Q{i+1}. {q_text}...</span>
                    <span style="font-size:13px;font-weight:700;color:{color}">{score}/100</span>
                </div>
                """, unsafe_allow_html=True)


def _submit_answer(q_obj: dict, answer: str):
    """Evaluate answer and advance to next question."""
    with st.spinner("🤖 AI is evaluating your answer..."):
        evaluation = evaluate_answer(
            question=q_obj["question"],
            ideal_answer=q_obj["ideal_answer"],
            key_points=q_obj["key_points"],
            user_answer=answer,
            section=q_obj["section"],
            difficulty=q_obj["difficulty"],
        )

    st.session_state.interview_session_results.append({
        "question_obj": q_obj,
        "user_answer": answer,
        "evaluation": evaluation,
    })

    # Show quick score feedback
    score = evaluation.get("score", 0)
    if score >= 75:
        st.success(f"🎉 Great answer! Score: {score}/100")
    elif score >= 50:
        st.warning(f"👍 Good attempt. Score: {score}/100")
    else:
        st.error(f"📚 Needs improvement. Score: {score}/100")

    # Advance
    st.session_state.interview_current_idx += 1
    st.session_state.voice_transcript_buffer = ""

    total = len(st.session_state.interview_questions_flat)
    if st.session_state.interview_current_idx >= total:
        _wrap_up_session()
    else:
        st.rerun()


def _wrap_up_session(early: bool = False):
    """Generate report and move to report phase."""
    results = st.session_state.interview_session_results
    if not results:
        st.warning("No answers recorded. Please answer at least one question.")
        return

    if early:
        st.info(f"Session ended early. {len(results)} questions answered.")

    with st.spinner("📊 Generating your feedback report..."):
        report = generate_feedback_report(results, st.session_state.interview_duration)
        st.session_state.interview_report = report
        st.session_state.interview_phase = "report"

    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 3: Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def _phase_report():
    report = st.session_state.interview_report
    if not report:
        st.error("No report available.")
        return

    score = report["overall_score"]
    band = report["performance_band"]
    band_color_map = {"Excellent": "#2ecc71", "Good": "#3498db", "Average": "#e67e22", "Needs Improvement": "#e94560"}
    band_color = band_color_map.get(band, "#888")

    # ── Score card ────────────────────────────────────────────────────────────
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:16px;padding:28px;margin-bottom:20px;text-align:center">
        <h1 style="color:#e94560;margin:0;font-size:52px;font-weight:800">{score}<span style="font-size:24px">/100</span></h1>
        <div style="color:{band_color};font-size:22px;font-weight:700;margin:6px 0">{band}</div>
        <p style="color:#a0b4d6;margin:0">Your Interview Performance Score</p>
    </div>
    """, unsafe_allow_html=True)

    # Section scores
    mc1, mc2, mc3, mc4 = st.columns(4)
    with mc1:
        st.metric("🎭 Behavioral", f"{report['behavioral_score']}/100")
    with mc2:
        st.metric("⚙️ Technical", f"{report['technical_score']}/100")
    with mc3:
        st.metric("✅ Questions Done", report["total_questions"])
    with mc4:
        st.metric("✅ Answered Well", len(report["well_answered"]))

    st.markdown("---")

    # ── Overall Summary ───────────────────────────────────────────────────────
    st.markdown("### 📋 Overall Assessment")
    st.info(report.get("overall_summary", ""))

    col_s, col_w = st.columns(2)
    with col_s:
        st.markdown("#### 💪 Strengths")
        for s in report.get("key_strengths", []):
            st.markdown(f"✅ {s}")

    with col_w:
        st.markdown("#### ⚠️ Areas to Improve")
        for w in report.get("weak_areas", []):
            st.markdown(f"🔸 {w}")

    # ── Section feedback ──────────────────────────────────────────────────────
    col_b, col_t = st.columns(2)
    with col_b:
        st.markdown("#### 🎭 Behavioral Feedback")
        st.markdown(report.get("behavioral_feedback", ""))
    with col_t:
        st.markdown("#### ⚙️ Technical Feedback")
        st.markdown(report.get("technical_feedback", ""))

    st.markdown("---")

    # ── Per-question results ──────────────────────────────────────────────────
    st.markdown("### 📊 Question-by-Question Results")

    for i, res in enumerate(report.get("session_results", [])):
        q_obj = res.get("question_obj", {})
        ev = res.get("evaluation", {})
        user_ans = res.get("user_answer", "")
        s = ev.get("score", 0)
        card_color = "#d4edda" if s >= 70 else "#fff3cd" if s >= 50 else "#f8d7da"
        border_color = "#28a745" if s >= 70 else "#ffc107" if s >= 50 else "#dc3545"

        with st.expander(f"Q{i+1}. {q_obj.get('question','')[:80]}... — {s}/100", expanded=False):
            st.markdown(f"""
            <div style="border-left:4px solid {border_color};background:{card_color};
                        border-radius:0 10px 10px 0;padding:14px;margin-bottom:12px">
                <b>{q_obj.get('section','').title()} | {q_obj.get('difficulty','')}</b><br>
                {q_obj.get('question','')}
            </div>
            """, unsafe_allow_html=True)

            tc1, tc2, tc3, tc4 = st.columns(4)
            tc1.metric("Overall", f"{ev.get('score',0)}/100")
            tc2.metric("Meaning Match", f"{ev.get('meaning_match',0)}/100")
            tc3.metric("Keyword Coverage", f"{ev.get('keyword_coverage',0)}/100")
            tc4.metric("Structure", f"{ev.get('structure_score',0)}/100")

            if user_ans:
                st.markdown("**Your Answer:**")
                st.info(user_ans)
            else:
                st.warning("*Question skipped*")

            kc1, kc2 = st.columns(2)
            with kc1:
                st.markdown("**✅ Keywords Covered:**")
                covered = ev.get("keywords_covered", [])
                st.markdown(", ".join(covered) if covered else "_None_")
            with kc2:
                st.markdown("**❌ Keywords Missed:**")
                missed = ev.get("keywords_missed", [])
                st.markdown(", ".join(missed) if missed else "_None — great!_")

            st.markdown("**💡 Coach Feedback:**")
            st.markdown(ev.get("brief_feedback", ""))

            if ev.get("strengths"):
                st.markdown("**Strengths:**")
                for s_item in ev["strengths"]:
                    st.markdown(f"✔ {s_item}")

            if ev.get("improvements"):
                st.markdown("**Improvements:**")
                for imp in ev["improvements"]:
                    st.markdown(f"• {imp}")

            if ev.get("improved_answer"):
                st.markdown("**📝 Suggested Improved Answer:**")
                st.success(ev["improved_answer"])

    st.markdown("---")

    # ── Keywords summary ──────────────────────────────────────────────────────
    st.markdown("### 🔑 Keyword Summary")
    kw_col1, kw_col2 = st.columns(2)
    with kw_col1:
        st.markdown("**Keywords Covered:**")
        covered_kws = report.get("keywords_covered", [])
        if covered_kws:
            badges = " ".join([f'<span style="background:#d4edda;color:#155724;padding:3px 10px;border-radius:15px;font-size:12px;margin:2px;display:inline-block">{k}</span>' for k in covered_kws])
            st.markdown(badges, unsafe_allow_html=True)
        else:
            st.markdown("_None_")
    with kw_col2:
        st.markdown("**Keywords Missed:**")
        missed_kws = report.get("keywords_missed", [])
        if missed_kws:
            badges = " ".join([f'<span style="background:#f8d7da;color:#721c24;padding:3px 10px;border-radius:15px;font-size:12px;margin:2px;display:inline-block">{k}</span>' for k in missed_kws])
            st.markdown(badges, unsafe_allow_html=True)
        else:
            st.markdown("_None — excellent coverage!_ 🎉")

    st.markdown("---")

    # ── Recommendations ───────────────────────────────────────────────────────
    st.markdown("### 🎯 Recommendations For Further Practice")
    for rec in report.get("recommendations", []):
        st.markdown(f"▶ {rec}")

    st.markdown("### 🚀 Next Steps")
    st.info(report.get("next_steps", ""))

    st.markdown("---")

    # ── Downloads ─────────────────────────────────────────────────────────────
    st.markdown("### 📥 Download Your Feedback Report")
    try:
        with st.spinner("Preparing report files..."):
            pdf_buf, docx_buf = export_feedback_report(report)

        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("📄 Download PDF Report", data=pdf_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                               mime="application/pdf", key="dl_report_pdf", use_container_width=True)
        with dl2:
            st.download_button("📝 Download DOCX Report", data=docx_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="dl_report_docx", use_container_width=True)
    except Exception as e:
        st.error(f"Report export failed: {e}")

    st.markdown("---")

    # ── Restart ───────────────────────────────────────────────────────────────
    rc1, rc2 = st.columns(2)
    with rc1:
        if st.button("🔁 Practice Again (New Session)", type="primary", use_container_width=True):
            _reset_interview_session()
            st.rerun()
    with rc2:
        if st.button("⚙️ Change Settings", use_container_width=True):
            st.session_state.interview_phase = "setup"
            st.rerun()


def _reset_interview_session():
    keys = [
        "interview_phase", "interview_qa_bank", "interview_questions_flat",
        "interview_current_idx", "interview_session_results", "interview_report",
        "voice_transcript_buffer",
    ]
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]
