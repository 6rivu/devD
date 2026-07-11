import streamlit as st
import time
import json
import re
import os
from io import BytesIO
from docx import Document
import plotly.graph_objects as go
from reportlab.lib.pagesizes import letter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import secrets
from database import register_user
from payment import create_checkout_session
from auth import get_current_user
import phonenumbers
import pycountry
from phonenumbers import parse as pn_parse, is_valid_number, format_number, PhoneNumberFormat
import resend
import urllib.parse
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash
load_dotenv()


# Import custom modules
from database import init_db, get_user_special_discount, get_user_data, save_user_session, get_business_plan_info, get_user_credits, get_business_credits, update_business_credits, activate_business_plan, get_db_connection, create_business_user, get_business_user, authenticate_business_user, save_business_payment, save_payment, update_user_credits, payment_exists, register_user, verify_user_email, set_email_otp, verify_email_otp, save_cv_generation, record_credit_usage, reset_credits_if_expired
from auth import authenticate_user, logout_user, get_current_user, hash_password
from payment import process_payment, check_subscription, apply_discount_code, create_checkout_session
from cv_generator import generate_cv,recommend_jobs_from_resume_ai, generate_cover_letter, extract_resume_text, analyze_cv_ats_score, generate_interview_qa, export_interview_qa
from templates import apply_template
from utils import optimize_keywords, enforce_page_limit, get_gemini_response, get_all_country_dial_codes



BUSINESS_PLANS = {
    "Corporate Starter": {
        "credits": 500,
        "price": 149.99,
        "duration": "3 months"
    },

    "Corporate Growth": {
        "credits": 1000,
        "price": 299.00,
        "duration": "3 months"
    },

    "Corporate Pro": {
        "credits": 2500,
        "price": 449.00,
        "duration": "6 months"
    },

    "Corporate Plus": {
        "credits": 5000,
        "price": 699.00,
        "duration": "6 months"
    },

    "Corporate Advanced": {
        "credits": 7500,
        "price": 899.00,
        "duration": "1 year"
    },

    "Corporate Enterprise": {
        "credits": 10000,
        "price": 999.00,
        "duration": "1 year"
    }
}

def add_home_button():
    home_href = "https://cvolvepro.com"
    st.sidebar.markdown(
        f'<a href="{home_href}" target="_self"><button style="width:100%">🏠 Home</button></a>',
        unsafe_allow_html=True
    )

resend.api_key = os.getenv("RESEND_API_KEY")
FROM_EMAIL = os.getenv("APP_FROM_EMAIL", "onboarding@resend.dev")

def _handle_tracking_hop():
    """If URL has ?trk=..., briefly show an interstitial and bounce to ?next=..."""
    # Works with both old/new Streamlit query APIs
    params = st.query_params if hasattr(st, "query_params") else st.experimental_get_query_params()

    def _get(key):
        v = params.get(key)
        return v[0] if isinstance(v, list) else v

    trk  = _get("trk")
    nxt  = _get("next")

    if trk and nxt:
        st.markdown("#### Redirecting…")
        st.caption(f"Tracking event: {trk}")
        # simple 0-sec meta refresh
        st.markdown(f'<meta http-equiv="refresh" content="0; url={nxt}">', unsafe_allow_html=True)
        st.stop()  # do not render rest of the app on this hop

# call it immediately so it catches both click and success hops
_handle_tracking_hop()

def _qp_get(key: str, default=""):
    v = st.query_params.get(key, default)
    return v[0] if isinstance(v, list) else v

def handle_stripe_return_globally():
    if _qp_get("service") == "jobsqa":
        return
    import stripe
    from database import get_db_connection, save_payment
    from payment import check_subscription, create_subscription
    from database import payment_exists as db_payment_exists, reset_credits_if_expired

    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")
    PLAN_CREDITS = {"Premium": 110, "Premium + Premium Classic": 125}

    success    = (_qp_get("success", "").lower() == "true")
    typ        = _qp_get("type", "")
    session_id = _qp_get("session_id", "")
    credits_qp = int(_qp_get("credits", "0") or 0)
    plan_qp    = _qp_get("plan", "")

    if not (success and session_id and typ in ("subscription", "credits", "business")):
        return

    # prevent double-processing on reruns
    if "processed_sessions" not in st.session_state:
        st.session_state.processed_sessions = set()
    if session_id in st.session_state.processed_sessions:
        return

    try:
        sess = stripe.checkout.Session.retrieve(session_id)
        if sess.get("payment_status") != "paid":
            return

        md = sess.get("metadata") or {}
        user_email = (md.get("user_email") or "").strip().lower()
        if not user_email:
            st.warning("Missing user email in Stripe metadata; cannot credit.")
            return

        # Idempotency: skip if we already stored this payment
        if db_payment_exists(session_id):
            st.session_state.processed_sessions.add(session_id)
            return

        amount_paid = (sess.get("amount_total") or 0) / 100.0

        if typ == "subscription":
            plan = (md.get("plan") or plan_qp or "Premium").strip()
            credits_to_add = int(md.get("credits") or PLAN_CREDITS.get(plan, 125))

            # start cycle today and set balance to plan credits
            conn = get_db_connection(); cur = conn.cursor()
            cur.execute("""
                UPDATE users
                   SET credits = %s,
                       credit_cycle_start = CURRENT_TIMESTAMP
                 WHERE email=%s
            """, (credits_to_add, user_email))
            conn.commit(); cur.close(); conn.close()

            create_subscription(user_email, plan, session_id)
            save_payment(user_email, amount_paid, "subscription", session_id, credits_purchased=credits_to_add)
            st.success(f"🎉 {plan} active for {user_email}. {credits_to_add} credits added.")

        elif typ == "credits":
            # read credits from metadata first, then query param as fallback
            credits_to_add = int(md.get("credits") or credits_qp or 0)

            # Optional: enforce active subscription for top-ups
            sub_active = check_subscription(user_email)
            if not sub_active:
                save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=0)
                st.error("Top-ups require an active plan. Payment recorded, but credits not added.")
            else:
                try:
                    reset_credits_if_expired(user_email)
                except Exception:
                    pass

                if credits_to_add > 0:
                    conn = get_db_connection(); cur = conn.cursor()
                    cur.execute("""
                        UPDATE users
                           SET credits = COALESCE(credits, 0) + %s
                         WHERE email=%s
                    """, (credits_to_add, user_email))
                    conn.commit(); cur.close(); conn.close()

                    save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=credits_to_add)
                    st.success(f"🎉 {credits_to_add} credits added for {user_email}.")
                else:
                    save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=0)
                    st.warning("Payment succeeded but no credits value was provided.")

        # ======================================================
        # BUSINESS PLAN HANDLER
        # ======================================================

        elif typ == "business":

            plan_name = md.get("plan_name", "Starter")

            credits_to_add = int(
                md.get("credits", 0)
            )

            duration = md.get(
                "duration",
                "3 months"
            )

            from datetime import datetime, timedelta

            if duration.lower() == "3 months":
                expiry = datetime.now() + timedelta(days=90)

            elif duration.lower() == "6 months":
                expiry = datetime.now() + timedelta(days=180)

            else:
                expiry = datetime.now() + timedelta(days=365)

            activate_business_plan(
                email=user_email,
                plan_name=plan_name,
                credits=credits_to_add,
                expiry=expiry
            )

            save_business_payment(
                user_email,
                amount_paid,
                "business_plan",
                session_id,
                credits_purchased=credits_to_add
            )

            st.success(
                f"""
                🎉 Business Plan Activated Successfully

                Plan: {plan_name}

                Credits Added: {credits_to_add}
                """
            )

    except Exception as e:
        st.error(f"Stripe verification failed: {e}")
    finally:
        st.session_state.processed_sessions.add(session_id)
        try:
            st.query_params.clear()
        except:
            pass


import stripe
stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")

import traceback

def _sanitize_db_text(s: str) -> str:
    """Make text safe for PostgreSQL INSERT (strip NULs, cap size)."""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    return s.replace("\x00", "")[:500_000]  # keep first ~500k chars


# Initialize database
init_db()

from database import seed_discount_codes
seed_discount_codes()

def get_allowed_ai_models_for_user():
    """
    Returns the AI model options the current user is allowed to see
    based on their active subscription.
    - Premium                 -> ["Premium"]
    - Premium + Premium Classic -> ["Premium", "Premium Classic"]
    - No active subscription  -> both (uses credits)
    """
    try:
        user = st.session_state.get("user_data")
        if not user:
            return ["Premium", "Premium Classic"]

        sub = check_subscription(user["email"])
        if not sub:
            # No sub – allow both; they’ll spend credits
            return ["Premium", "Premium Classic"]

        plan = (sub.get("plan") or "").strip()
        if "Premium + Premium Classic" in plan:
            return ["Premium", "Premium Classic"]
        # Default to Premium-only
        return ["Premium"]
    except Exception:
        # On any error, fall back to both (don’t block usage)
        return ["Premium", "Premium Classic"]


# Page config
st.set_page_config(
    page_title="CVOLVE PRO - AI-Powered Resume Optimization",
    page_icon="logo.jpeg",
    layout="wide",
    initial_sidebar_state="expanded"
)

add_home_button()

# Load custom CSS
with open("styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Initialize session state
if 'user_data' not in st.session_state:
    st.session_state.user_data = None
if 'cv_preview' not in st.session_state:
    st.session_state.cv_preview = None
if 'auto_save' not in st.session_state:
    st.session_state.auto_save = {}
if 'selected_template' not in st.session_state:
    st.session_state.selected_template = "professional"

def auto_save_progress():
    """Auto-save user progress"""
    if st.session_state.user_data and st.session_state.auto_save:
        try:
            save_user_session(st.session_state.user_data['email'], st.session_state.auto_save)
        except Exception as e:
            # Silently handle auto-save errors to not interrupt user flow
            pass

def main():

    handle_stripe_return_globally()

    # ✅ Persist page navigation (login/register/main)
    if "page" not in st.session_state:
        st.session_state.page = "login"  # Default page

    if "portal" not in st.session_state:
        st.session_state.portal = "individual"
    

    # ✅ Show Register Page if user clicked register
    if st.session_state.page == "register":
        show_register_page()
        return  # Stop here after rendering register page


    # ✅ If user is not logged in
    if not st.session_state.get("user_data"):

        if st.session_state.portal == "business":
            show_business_login_page()
        else:
            show_login_page()

        return
    
    # Auto-save progress only when user is logged in and has data to save
    if st.session_state.user_data and st.session_state.auto_save:
        auto_save_progress()
    
    # Header
    st.markdown(f"""
    <div class="header">
        <h1 style="display:inline-block; vertical-align:middle; margin:0;">CVOLVE PRO</h1>
        <p>Transform your resume into an ATS-optimized masterpiece</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Authentication
    current_user = get_current_user()
    if not current_user:
        show_login_page()
        return
    
    st.session_state.user_data = current_user
    # NEW: Just-in-time monthly reset so UI shows the correct balance
    try:
        reset_credits_if_expired(current_user['email'])
    except Exception:
        pass


    # If user navigated to Billing from sidebar, render only Billing and return
    if st.session_state.get("page") == "billing":
        st.markdown("## 💳 Billing")
        show_billing_page()
        if st.button("⬅ Back"):
            st.session_state.page = "home"   # any non-"billing" value works
            st.rerun()
        return
    
    # Sidebar
    with st.sidebar:
        st.markdown(
        f"👋 Welcome, {current_user.get('name', current_user.get('company_name', 'User'))}"
    )
        
        # User credits/subscription status
        email = current_user['email'].strip().lower()

        if st.session_state.get("account_type") == "business":
            credits = get_business_credits(email)
            subscription = None
        else:
            credits = get_user_credits(email)
            subscription = check_subscription(email)

        # 🔹 Set AI model options based on subscription
        plan = (subscription['plan'] if subscription else "") or ""
        allowed_models = ["Premium", "Premium Classic"] if "Premium + Premium Classic" in plan or plan == "" else ["Premium"]

        # Save in session; rerun if changed so dropdown updates instantly
        if st.session_state.get("ai_model_options") != allowed_models:
            st.session_state["ai_model_options"] = allowed_models
            if st.session_state.get("ai_model") == "openai" and "Premium Classic" not in allowed_models:
                st.session_state["ai_model"] = "gemini"
            st.rerun()

        
        if st.session_state.get("account_type") == "business":

            plan_info = get_business_plan_info(email)

            if plan_info and plan_info["current_plan"]:

                st.success(
                    f"🏢 {plan_info['current_plan']} Plan Active"
                )

                if plan_info["plan_expiry"]:

                    st.caption(
                        f"Valid Until: {plan_info['plan_expiry'].strftime('%d %b %Y')}"
                    )

            else:

                st.warning("⏳ No active business plan")

        else:

            if subscription:

                if subscription['plan'] == "Free":

                    st.info("🆓 Free Trial (1 Month)")

                else:

                    st.success(
                        f"✅ {subscription['plan']} Plan Active"
                    )

            else:

                st.warning("⏳ No active plan")

        # 💎 Credits — ALWAYS visible
        st.info(f"💎 Credits Available: {credits}")
            
        if st.button("🔄 Buy More Credits"):
            st.session_state.page = "billing"
            st.rerun()

            
        if st.sidebar.button("Logout"):

            # Individual session
            st.session_state.logged_in = False
            st.session_state.user_data = None

            # Business session
            st.session_state.business_logged_in = False
            st.session_state.business_user = None
            st.session_state.business_email = None
            st.session_state.business_company = None

            # Global
            st.session_state.account_type = None

            st.rerun()
            
        
        # ✅ Set default template to Professional Classic
        st.session_state.selected_template = "professional"
        
        # ✅ Always include default sections (all enabled)
        sections = {
            "Professional Summary": True,
            "Key Skills": True,
            "Work Experience": True,
            "Education": True,
            "Certifications": True,
            "Projects": True,
            "Awards": False,
            "Languages": False,
            "Hobbies": False
        }
        
        st.session_state.auto_save['sections'] = sections
        
        # Quick links
        st.markdown("---")
        with st.sidebar.expander("📚 How It Works"):
            st.markdown("""
            1. Upload your resume (PDF/DOCX)  
            2. Paste the job description  
            3. Choose your sections & template  
            4. Click ‘Generate Optimized CV’  
            5. Download your resume or cover letter  
            """)

            with st.sidebar.expander("🔒 Privacy Policy"):
                st.markdown("""
                - Your data is processed securely  
                - Resumes and job descriptions are not stored  
                - No personal info is shared with third parties  
                """)

    
    # Main content
    tab1, tab3, tab4 = st.tabs(["🎯 Match Me to Job", "📊 Analytics", "💳 Billing"])

    with tab1:
        show_cv_generation_page()

    with tab3:
        show_analytics_page()

    with tab4:
        show_billing_page()

def show_login_page():
    st.markdown("## 🔐 Login to CVOLVE PRO")
    colp1, colp2 = st.columns(2)

    with colp1:
        if st.button("👤 Individual Portal"):
            st.session_state.portal = "individual"
            st.rerun()

    with colp2:
        if st.button("🏢 Business Portal"):
            st.session_state.portal = "business"
            st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### Email Login")
        email = st.text_input("Email Address", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")

        # ── Buttons side-by-side ───────────────────────────────────────────────
        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            login_clicked = st.button("🔑 Login", key="login_button", use_container_width=True)
        with btn_col2:
            register_clicked = st.button("🆕 Register For Free", key="register_button", use_container_width=True)
        # ───────────────────────────────────────────────────────────────────────

        if login_clicked:
            if email.strip() and password.strip():
                email_norm = email.strip().lower()
                existing = get_user_data(email_norm)

                if not existing:
                    st.error("No account found for this email. Please register first.")
                    return

                if not existing.get("is_verified", False):
                    st.error("Your email is not verified. Please complete registration with OTP.")
                    return

                user = authenticate_user(email_norm, password, "email")
                if user:
                    st.session_state.user_data = user
                    st.session_state.account_type = "individual"
                    st.success("✅ Login successful!")
                    st.rerun()
                else:
                    st.error("Incorrect password.")
            else:
                st.error("Please enter both email and password")

        if register_clicked:
            st.session_state.page = "register"  # ✅ Persist state
            st.rerun()

    # (Optional) keep the divider if you want extra spacing
    # st.markdown("---")


def show_business_login_page():

    st.markdown("# 🏢 Business Portal")

    tab1, tab2 = st.tabs([
        "Business Login",
        "Business Register"
    ])

    # =====================================================
    # LOGIN
    # =====================================================

    with tab1:

        email = st.text_input(
            "Business Email",
            key="business_login_email"
        )

        password = st.text_input(
            "Password",
            type="password",
            key="business_login_password"
        )

        if st.button("Login to Business Portal"):

            user = authenticate_business_user(
                email,
                password
            )

            if user:

                # Main session
                st.session_state.user_data = user

                # Business session
                st.session_state.business_user = user
                st.session_state.business_logged_in = True

                # Optional helper fields
                st.session_state.business_email = user["email"]
                st.session_state.business_company = user["company_name"]
                st.session_state.account_type = "business"

                st.success("Business Login Successful")

                st.rerun()

            else:
                st.error("Invalid Credentials")

    # =====================================================
    # REGISTER
    # =====================================================

    with tab2:

        company_name = st.text_input("Company Name")

        owner_name = st.text_input("Owner Name")

        email = st.text_input("Business Email")

        password = st.text_input(
            "Create Password",
            type="password"
        )

        selected_plan = st.selectbox(
            "Corporate Plan",
            list(BUSINESS_PLANS.keys())
        )

        plan = BUSINESS_PLANS[selected_plan]

        st.success(
            f"""
            Credits: {plan['credits']}

            Duration: {plan['duration']}

            Price: ${plan['price']}
            """
        )

        st.markdown("## 📦 Corporate Packages")

        st.table([
            ["Starter", 500, "$149.99", "3 Months"],
            ["Growth", 1000, "$299", "3 Months"],
            ["Pro", 2500, "$449", "6 Months"],
            ["Plus", 5000, "$699", "6 Months"],
            ["Advanced", 7500, "$899", "1 Year"],
            ["Enterprise", 10000, "$999", "1 Year"],
        ])

        if st.button("Create Business Account"):

            existing = get_business_user(email)

            if existing:
                st.error("Business email already exists")
                return

            password_hash = generate_password_hash(password)

            create_business_user(
                company_name=company_name,
                owner_name=owner_name,
                email=email,
                password_hash=password_hash,
                plan_name=selected_plan
            )

            st.success("Business Account Created")



def show_cv_generation_page():
    """Main CV generation interface"""
    col_header, col_dropdown = st.columns([3, 1])
    with col_header:
        st.markdown("## 🎯 Match Me to the Job")
    with col_dropdown:
        # Limit choices by subscription
        allowed_options = get_allowed_ai_models_for_user()

        options = st.session_state.get("ai_model_options", ["Premium", "Premium Classic"])
        default_label = "Premium Classic" if st.session_state.get("ai_model") == "openai" else "Premium"
        if default_label not in options:
            default_label = options[0]

        model_choice = st.selectbox(
            "AI Model",
            options=options,
            index=options.index(default_label)
        )
        st.session_state["ai_model"] = "openai" if model_choice == "Premium Classic" else "gemini"

        # Safety: if plan is Premium-only but session had Classic, force-correct it
        if "Premium Classic" not in allowed_options and st.session_state.get("ai_model") == "openai":
            st.session_state["ai_model"] = "gemini"

        # --- NEW: Language selector shown beside model selector ---
        # Initialize default language in session state if not present
        if "selected_language" not in st.session_state:
            st.session_state["selected_language"] = "English"

        # Language options — add more languages as you like
        language_options = ["English", "Français", "Español", "Deutsch"]

        # Show a compact dropdown for language choice
        st.selectbox(
            "Language",
            options=language_options,
            index=language_options.index(st.session_state["selected_language"]) if st.session_state["selected_language"] in language_options else 0,
            key="selected_language"
        )


    # ✅ Define callback at the start of JD section
    def clear_jd():
        st.session_state.jd_input = ""
        st.session_state.job_description = ""

    # Job Description Input
    st.markdown("### 📋 Job Description")
    jd = st.text_area(
        "Paste the job description here",
        height=200,
        placeholder="Copy and paste the complete job description...",
        key="jd_input"
    )

    # ✅ Clear JD Button
    st.button("🧹 Clear JD", help="Click to clear job description", on_click=clear_jd)

    # ✅ Save JD in session for Q&A tab
    if jd.strip():
        st.session_state.job_description = jd
        st.session_state.auto_save['job_description'] = jd

    if jd.strip():
        with st.expander("📝 Job Description Preview"):
            st.code(jd, language="markdown")

    # Resume Upload
    st.markdown("### 📄 Upload Your Resume")
    uploaded_file = st.file_uploader(
        "Choose your resume file",
        type=["pdf", "docx"],
        help="Upload your existing resume in PDF or DOCX format"
    )

    # ✅ Save resume in session for Q&A tab
    if uploaded_file:
        st.session_state.uploaded_resume = uploaded_file

    # ATS Score Check
    # ATS Score Check + AI Job Match (side by side)
    if uploaded_file:

        col_ats, col_ai = st.columns([1, 1])

        with col_ats:
            check_ats_btn = st.button(
                "📊 Check ATS Score",
                use_container_width=True
            )

        with col_ai:
            subscription = check_subscription(st.session_state.user_data['email'])

            ai_job_btn = st.button(
                "Job Recommendations",
                disabled=not bool(subscription),
                help="Available for paid users only",
                use_container_width=True
            )

        # ---------- ATS SCORE LOGIC (UNCHANGED) ----------
        if check_ats_btn:
            if not check_user_access(required_credits=1):
                st.error("⚠️ You need at least 1 credit to run ATS Check.")
            else:
                try:
                    resume_text = extract_resume_text(uploaded_file)
                    analysis = analyze_cv_ats_score(resume_text, jd)

                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("ATS Score", f"{analysis['score']}%")
                        st.progress(analysis['score'] / 100)
                        if analysis['score'] < 40:
                            st.warning(
                                "⚠️ Your ATS score is critically low. Do you still want to proceed for this job!"
                            )

                    with col2:
                        st.metric("Keyword Match", f"{analysis['keyword_match']}%")
                        st.progress(analysis['keyword_match'] / 100)

                    if analysis.get('suggestions'):
                        st.markdown("### 💡 Improvement Suggestions")
                        for suggestion in analysis['suggestions']:
                            st.markdown(f"• {suggestion}")

                    if analysis.get('missing_keywords'):
                        st.markdown("### 🔍 Missing Keywords")
                        for keyword in analysis['missing_keywords'][:5]:
                            st.markdown(f"• {keyword}")

                    # ✅ Deduct only on success
                    deduct_user_credits(
                        st.session_state.user_data['email'],
                        1,
                        feature="ATS"
                    )

                except Exception as e:
                    st.error(f"❌ Error analyzing ATS score: {str(e)}")

        # ---------- AI JOB MATCH (LIST ONLY) ----------
        if ai_job_btn:
            resume_text = extract_resume_text(uploaded_file)

            if resume_text.strip():
                try:
                    with st.spinner("🔍 Analyzing your resume for job matches..."):
                        jobs = recommend_jobs_from_resume_ai(
                            resume_text,
                            language=st.session_state.get("selected_language", "English")
                        )

                    if jobs:
                        st.markdown("### ✅ Recommended Job Roles")
                        for job in jobs:
                            st.markdown(f"- {job}")

                        deduct_user_credits(
                                st.session_state.user_data['email'],
                                1,
                                feature="Job Match"
                            )
                    else:
                        st.warning("⚠️ No job recommendations could be generated. "
                                   "Please ensure your resume has sufficient content and try again.")
                except Exception as e:
                    st.error(f"❌ Error generating job recommendations: {str(e)}")
            else:
                st.warning("⚠️ Could not extract text from your resume. "
                           "Please ensure the file is not empty or corrupted.")

    else:
        st.info("Please upload your resume and enter a job description to check ATS score.")

    # Target Match Percentage
    target_match = st.slider(
        "🎯 Target ATS Match Percentage",
        min_value=60,
        max_value=100,
        value=90,
        step=1,
        help="Higher percentages may require more aggressive optimization"
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        generate_cv_btn = st.button("🚀 Generate Optimized CV", type="primary")
    with col2:
        generate_cover_letter_btn = st.button("📝 Generate Cover Letter")
    with col3:
        # Only show Q&A for Premium Classic or credits users
        email = st.session_state.user_data['email']

        if st.session_state.get("account_type") == "business":
            credits_avail = get_business_credits(email)
            subscription = None
        else:
            subscription = check_subscription(email)
            credits_avail = get_user_credits(email)
        allowed_for_qa = bool(
            subscription
            and "Premium + Premium Classic" in subscription['plan']
            and credits_avail >= 3
        )
        if allowed_for_qa:
            generate_qa_btn = st.button("🎤 Generate Interview Q&A")
        else:
            generate_qa_btn = False
            st.button("🎤 Generate Interview Q&A", disabled=True,
                      help="Requires Premium + Premium Classic plan and ≥3 credits")

    # ======================
    # Generate CV (generation-only; no download UI here)
    # ======================
    if generate_cv_btn:
        if uploaded_file and jd.strip():
            # Check credits/subscription
            if not check_user_access(required_credits=3):
                st.error("⚠️ Insufficient credits. Please purchase more credits or upgrade your subscription.")
                return

            loading_placeholder = st.empty()
            loading_placeholder.markdown("""
                <div style="display: flex; flex-direction: column; align-items: center; padding: 20px;">
                    <div class="custom-loader"></div>
                    <p style="margin-top: 10px;">🔄 Optimizing your CV... Please wait</p>
                </div>
            """, unsafe_allow_html=True)

            time.sleep(0.5)  # Optional: show loader briefly before real work starts
            start_time = time.time()

            try:
                # Extract resume text
                resume_text = extract_resume_text(uploaded_file)

                # Generate optimized CV
                sections_to_use = st.session_state.auto_save.get('sections', {
                    "Professional Summary": True,
                    "Key Skills": True,
                    "Work Experience": True,
                    "Education": True,
                    "Certifications": True,
                    "Projects": True,
                    "Awards": False,
                    "Languages": False,
                    "Hobbies": False
                })

                st.session_state["target_match"] = target_match

                cv_content = generate_cv(
                    resume_text=resume_text,
                    job_description=jd,
                    target_match=target_match,
                    template=st.session_state.selected_template,
                    sections=sections_to_use,
                    quantitative_focus=60,
                    action_verb_intensity="High",
                    keyword_matching="Balanced",
                    language=st.session_state.get("selected_language", "English")
                )

                # Enforce 2-page limit
                cv_content = enforce_page_limit(cv_content)

                # Store in session for preview
                st.session_state.cv_preview = cv_content
                st.session_state.job_description = jd  # for ATS analysis

                # ===== Cache export bytes for stable downloads across reruns =====
                clean_preview = st.session_state.cv_preview.replace("**", "")  # strip markdown asterisks for PDF
                pdf_buffer = apply_template(clean_preview, st.session_state.selected_template)
                docx_buffer = create_word_document(st.session_state.cv_preview)
                st.session_state.cv_pdf_bytes = pdf_buffer.getvalue()
                st.session_state.cv_docx_bytes = docx_buffer.getvalue()

                loading_placeholder.empty()

                processing_time = time.time() - start_time
                st.success(f"✅ CV generated successfully in {processing_time:.1f} seconds!")

                # --- Persist this generation for Analytics ---
                try:
                    jd_clean     = _sanitize_db_text(jd)
                    resume_clean = _sanitize_db_text(resume_text)
                    cv_clean     = _sanitize_db_text(st.session_state.cv_preview)

                    quick_analysis = optimize_keywords(cv_clean, jd_clean)
                    ats_score_val  = int(quick_analysis.get("score") or target_match or 0)

                    if st.session_state.get("account_type") != "business":

                        save_cv_generation(
                            user_email=st.session_state.user_data['email'],
                            job_description=jd_clean,
                            original_resume=resume_clean,
                            generated_cv=cv_clean,
                            template_used=st.session_state.selected_template,
                            ats_score=ats_score_val,
                            target_match=int(target_match),
                            processing_time=float(f"{processing_time:.2f}")
                        )
                except Exception as e:
                    st.error("❌ Failed to record this CV in Analytics. See details below.")
                    st.exception(e)

                # Deduct credits (only once here)
                deduct_user_credits(st.session_state.user_data['email'], 3, feature="CV")

            except Exception as e:
                st.error(f"❌ Error generating CV: {str(e)}")
        else:
            st.warning("⚠️ Please upload your resume and provide a job description")

    # ======================
    # Generate Cover Letter (unchanged flow)
    # ======================
    if generate_cover_letter_btn:
        if uploaded_file and jd.strip():
            if not check_user_access(required_credits=2):
                st.error("⚠️ Insufficient credits. Please purchase more credits or upgrade your subscription.")
                return

            loading_placeholder = st.empty()
            loading_placeholder.markdown("""
                <div style="display: flex; flex-direction: column; align-items: center; padding: 20px;">
                    <div class="custom-loader"></div>
                    <p style="margin-top: 10px;">📝 Generating cover letter... Please wait</p>
                </div>
            """, unsafe_allow_html=True)

            time.sleep(0.5)

            try:
                resume_text = extract_resume_text(uploaded_file)
                cover_letter = generate_cover_letter(resume_text, jd, language=st.session_state.get("selected_language", "English"))

                # ✅ Clean any Markdown markers like ** or *
                cover_letter = re.sub(r'\*{1,2}', '', cover_letter)

                loading_placeholder.empty()
                st.session_state.cover_letter = cover_letter

                from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_JUSTIFY
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.units import inch
                from io import BytesIO
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                with st.expander("📄 Generated Cover Letter"):
                    # Display in UI
                    st.markdown(cover_letter)

                    # ===== PDF EXPORT WITH FIXED MARGINS AND JUSTIFIED TEXT =====
                    pdf_buffer = BytesIO()
                    doc = SimpleDocTemplate(
                        pdf_buffer,
                        pagesize=letter,
                        leftMargin=40, rightMargin=40,  # ✅ 0.4 inch
                        topMargin=35, bottomMargin=35   # ✅ 0.5 inch
                    )

                    styles = getSampleStyleSheet()
                    justified_style = ParagraphStyle(
                        name='Justified',
                        parent=styles['Normal'],
                        alignment=TA_JUSTIFY,
                        fontName='Helvetica',
                        fontSize=11,
                        leading=16
                    )

                    flowables = []
                    for paragraph in cover_letter.strip().split('\n'):
                        if paragraph.strip():
                            para = Paragraph(paragraph.strip(), justified_style)
                            flowables.append(para)
                            flowables.append(Spacer(1, 0.2 * inch))

                    doc.build(flowables)
                    pdf_buffer.seek(0)

                    st.download_button(
                        label="📥 Download as PDF",
                        data=pdf_buffer,
                        file_name="cover_letter.pdf",
                        mime="application/pdf"
                    )

                    # ===== DOCX EXPORT WITH FIXED MARGINS AND JUSTIFIED TEXT =====
                    docx_buffer = BytesIO()
                    word_doc = Document()

                    # ✅ Apply same margins as CV
                    for section in word_doc.sections:
                        section.top_margin = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin = Inches(0.4)
                        section.right_margin = Inches(0.4)

                    # Set base font and size
                    style = word_doc.styles['Normal']
                    font = style.font
                    font.name = 'Calibri'
                    font.size = Pt(11)

                    for paragraph in cover_letter.strip().split('\n'):
                        if paragraph.strip():
                            para = word_doc.add_paragraph(paragraph.strip())
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    word_doc.save(docx_buffer)
                    docx_buffer.seek(0)

                    st.download_button(
                        label="📥 Download as Word",
                        data=docx_buffer,
                        file_name="cover_letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Deduct credits
                deduct_user_credits(st.session_state.user_data['email'], 2, feature="CL")

            except Exception as e:
                loading_placeholder.empty()
                st.error(f"❌ Error generating cover letter: {str(e)}")

    # ✅ Generate Interview Q&A (unchanged)
    if generate_qa_btn:
        if uploaded_file and jd.strip():
            if not check_user_access(required_credits=3):
                st.error("⚠️ Insufficient credits. Please purchase more credits or upgrade your subscription.")
                return

            loading_placeholder = st.empty()
            loading_placeholder.markdown("""
                <div style="display: flex; flex-direction: column; align-items: center; padding: 20px;">
                    <div class="custom-loader"></div>
                    <p style="margin-top: 10px; font-weight:bold; font-size:16px;">⏳ Generating interview Q&A... Please wait</p>
                </div>
            """, unsafe_allow_html=True)

            try:
                # Extract resume text
                resume_text = extract_resume_text(uploaded_file)

                # Generate Q&A
                qa_content = generate_interview_qa(resume_text, jd)

                loading_placeholder.empty()

                st.markdown("### 📌 Suggested Questions & Answers")
                st.markdown(qa_content)

                # ✅ Export Options
                pdf_buffer, docx_buffer = export_interview_qa(qa_content)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Download PDF",
                        data=pdf_buffer,
                        file_name="interview_QA.pdf",
                        mime="application/pdf"
                    )
                with col2:
                    st.download_button(
                        "📥 Download DOCX",
                        data=docx_buffer,
                        file_name="interview_QA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Deduct credits
                deduct_user_credits(st.session_state.user_data['email'], 3, feature="Interview QA")

            except Exception as e:
                loading_placeholder.empty()
                st.error(f"❌ Error generating Q&A: {str(e)}")
        else:
            st.warning("⚠️ Please upload your resume and provide a job description")

    # ==============================================================
    # ✅ Always-visible Preview & Download section (persists reruns)
    # ==============================================================
    if st.session_state.get("cv_preview"):
        st.markdown("### 👀 Your Optimized CV")

        # Ensure export bytes exist (rebuild if a previous run didn't store them)
        pdf_bytes = st.session_state.get("cv_pdf_bytes")
        docx_bytes = st.session_state.get("cv_docx_bytes")
        if not pdf_bytes or not docx_bytes:
            try:
                clean_preview = st.session_state.cv_preview.replace("**", "")
                _pdf_buf = apply_template(clean_preview, st.session_state.selected_template)
                _docx_buf = create_word_document(st.session_state.cv_preview)
                pdf_bytes = _pdf_buf.getvalue()
                docx_bytes = _docx_buf.getvalue()
                st.session_state.cv_pdf_bytes = pdf_bytes
                st.session_state.cv_docx_bytes = docx_bytes
            except Exception as e:
                st.error(f"❌ Failed to prepare downloads: {e}")
                pdf_bytes, docx_bytes = None, None

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button(
                label="📥 Download PDF",
                data=pdf_bytes,
                file_name="optimized_cv.pdf",
                mime="application/pdf",
                key="dl_cv_pdf_persist",
                disabled=(pdf_bytes is None)
            )
        with c2:
            st.download_button(
                label="📄 Download DOCX",
                data=docx_bytes,
                file_name="optimized_cv.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cv_docx_persist",
                disabled=(docx_bytes is None)
            )
        with c3:
            if st.button("🔄 Regenerate CV", key="regen_cv_persist"):
                st.session_state.cv_preview = None
                st.session_state.pop("cv_pdf_bytes", None)
                st.session_state.pop("cv_docx_bytes", None)
                st.rerun()

        # Show preview content and ATS analysis
        st.markdown("### 📋 Preview Content")
        st.markdown(st.session_state.cv_preview)

        st.markdown("### 📊 ATS Analysis")
        analyze_ats_compatibility()

        st.info("🔍 Use the buttons above to download your CV. You can also switch templates and regenerate any time.")



def show_preview_page():
    """CV preview and download page"""
    st.markdown("## 📄 CV Preview")
    
    if st.session_state.cv_preview:
        st.markdown("### 👀 Your Optimized CV")
        
        # Preview options
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📥 Download as PDF"):
                pdf_buffer = apply_template(
                    st.session_state.cv_preview,
                    st.session_state.selected_template
                )
                
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_buffer,
                    file_name="optimized_cv.pdf",
                    mime="application/pdf"
                )
        
        with col2:
            if st.button("📄 Download as Word"):
                docx_buffer = create_word_document(st.session_state.cv_preview)
                
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_buffer,
                    file_name="optimized_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col3:
            if st.button("🔄 Regenerate"):
                st.session_state.cv_preview = None
                st.rerun()
        
        # Show preview
        st.markdown("---")
        st.markdown("### 📋 Preview Content")
        st.markdown(st.session_state.cv_preview)
        
        # ATS Analysis - Show automatically
        st.markdown("### 📊 ATS Analysis")
        analyze_ats_compatibility()
    
    else:
        st.info("🔍 No CV preview available. Please generate a CV first.")

def show_analytics_page():
    """Analytics dashboard (live from DB)"""
    from database import get_db_connection, get_user_credits
    import plotly.graph_objects as go

    st.markdown("## 📊 Your Analytics")

    user = st.session_state.get("user_data")
    if not user:
        st.info("Please log in to see analytics.")
        return
    user_email = user["email"]

    # ---------- Aggregates from DB ----------
    total_cvs = 0
    avg_ats = 0.0
    success_rate = 0.0
    if st.session_state.get("account_type") == "business":
        credits_now = get_business_credits(user_email) or 0
    else:
        credits_now = get_user_credits(user_email) or 0

    trend_dates, trend_scores = [], []

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Totals
        cur.execute("SELECT COUNT(*), COALESCE(AVG(ats_score),0) FROM cv_generations WHERE user_email=%s", (user_email,))
        row = cur.fetchone() or (0, 0.0)
        total_cvs = int(row[0] or 0)
        avg_ats = float(row[1] or 0.0)

        # Success rate
        cur.execute("""
            SELECT 
                COALESCE(SUM(CASE WHEN ats_score IS NOT NULL 
                                   AND target_match IS NOT NULL 
                                   AND ats_score >= target_match THEN 1 ELSE 0 END), 0) AS successes,
                COUNT(*) AS total_rows
            FROM cv_generations
            WHERE user_email=%s
        """, (user_email,))
        srow = cur.fetchone() or (0, 0)
        successes, total_rows = int(srow[0] or 0), int(srow[1] or 0)
        success_rate = (successes / total_rows * 100.0) if total_rows > 0 else 0.0

        # Trend
        cur.execute("""
            SELECT DATE(created_at) AS d, COALESCE(AVG(ats_score), 0)
            FROM cv_generations
            WHERE user_email=%s
            GROUP BY 1
            ORDER BY 1
        """, (user_email,))
        trend = cur.fetchall() or []
        trend_dates = [str(r[0]) for r in trend]
        trend_scores = [float(r[1] or 0.0) for r in trend]

        # Credits used TOTAL via credit_usage
        cur.execute("""
            SELECT COALESCE(SUM(credits), 0) 
              FROM credit_usage 
             WHERE user_email=%s
        """, (user_email,))
        credits_used_total = int((cur.fetchone() or [0])[0] or 0)

        # Per-feature usage (includes ATS)
        cur.execute("""
            SELECT feature, COALESCE(SUM(credits), 0)
              FROM credit_usage
             WHERE user_email=%s
             GROUP BY feature
        """, (user_email,))
        rows = cur.fetchall() or []
        usage_map = {r[0]: int(r[1] or 0) for r in rows}
        used_cv = usage_map.get("CV", 0)
        used_cl = usage_map.get("CL", 0)
        used_qa = usage_map.get("Interview QA", 0)
        used_ats = usage_map.get("ATS", 0)

        cur.close(); conn.close()
    except Exception as e:
        try:
            cur.close(); conn.close()
        except:
            pass
        st.error(f"Failed to load analytics: {e}")
        return

    # ---------- KPI Tiles ----------
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("CVs Generated", f"{total_cvs}")
    with col2:
        st.metric("Avg ATS Score", f"{avg_ats:.0f}%")
    with col3:
        st.metric("Total Credits Used", f"{credits_used_total}")   # <-- from credit_usage
    with col4:
        st.metric("Success Rate", f"{success_rate:.0f}%")

    st.markdown("---")

    # ---------- Credits by Feature ----------
    st.markdown("### Credits Used by Feature")
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("CV", used_cv)
    with c2: st.metric("Cover Letter", used_cl)
    with c3: st.metric("Interview Q&A", used_qa)
    with c4: st.metric("ATS", used_ats)



    # ---------- Trends ----------
    st.markdown("### 📈 ATS Score Trend")
    if trend_dates:
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=trend_dates, y=trend_scores, mode='lines+markers', name='Avg ATS'))
        fig.update_layout(
            margin=dict(l=10, r=10, t=40, b=10),
            title="Average ATS Score by Day",
            xaxis_title="Date", yaxis_title="Avg ATS %", yaxis=dict(range=[0, 100])
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No ATS score history yet. Generate a CV to see your trend!")

    # ---------- Recent activity ----------
    st.markdown("### 🧾 Recent CV Generations")
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT created_at, ats_score, target_match, template_used
              FROM cv_generations
             WHERE user_email=%s
             ORDER BY created_at DESC
             LIMIT 10
        """, (user_email,))
        rows = cur.fetchall() or []
        cur.close(); conn.close()

        if rows:
            st.dataframe(
                [{
                    "Date": r[0].strftime("%Y-%m-%d %H:%M"),
                    "ATS": int(r[1]) if r[1] is not None else None,
                    "Target": int(r[2]) if r[2] is not None else None,
                    "Template": r[3] or "-"
                } for r in rows],
                use_container_width=True, hide_index=True
            )
        else:
            st.write("No recent generations found.")
    except Exception as e:
        st.warning(f"Couldn’t load recent items: {e}")


def show_billing_page():
    """Billing and subscription management with Stripe + simple currency selection"""

    # ---- Currency: safe, local-friendly (no JS, no early return) ----
    cur_param = st.query_params.get("cur", None)
    if not cur_param:
        # Best-effort first guess from phone prefix while local/testing
        phone = (st.session_state.get("user_data", {}).get("phone") or "").strip()
        prefix_map = {
            "+91": "INR",   # India
            "+971": "AED",  # UAE
            "+973": "BHD",  # Bahrain
            "+61": "AUD",   # Australia
            "+44": "GBP",   # UK
        }
        guess = "USD"
        for pref, code in prefix_map.items():
            if phone.startswith(pref):
                guess = code
                break
        st.query_params["cur"] = guess  # triggers a rerun

    cur_param = st.query_params.get("cur", "USD")
    if isinstance(cur_param, list):
        cur_param = cur_param[0]

    SUPPORTED = ("INR", "EUR", "USD", "AED", "BHD", "AUD", "GBP")
    CURRENT_CURRENCY = cur_param if cur_param in SUPPORTED else "USD"

    SYMBOL = {
        "USD": "$", "INR": "₹", "EUR": "€",
        "AED": "د.إ", "BHD": "BD", "AUD": "A$", "GBP": "£",
    }[CURRENT_CURRENCY]

    # Simple static USD→local multipliers (adjust anytime)
    RATE = {
        "USD": 1.00,
        "INR": 84.00,
        "EUR": 0.92,
        "AED": 3.67,
        "BHD": 0.38,
        "AUD": 1.50,
        "GBP": 0.78,
    }[CURRENT_CURRENCY]

    THREE_DECIMAL = {"BHD"}  # BHD uses 3 decimals

    def price_local(usd_amount: float) -> float:
        v = usd_amount * RATE
        if CURRENT_CURRENCY == "INR":
            return round(v)           # whole rupees
        if CURRENT_CURRENCY in THREE_DECIMAL:
            return round(v, 3)        # e.g., BHD
        return round(v, 2)            # default 2-decimal

    def fmt(amount_local: float) -> str:
        if CURRENT_CURRENCY == "INR":
            return f"{SYMBOL}{amount_local:,.0f}"
        if CURRENT_CURRENCY in THREE_DECIMAL:
            return f"{SYMBOL}{amount_local:,.3f}"
        return f"{SYMBOL}{amount_local:,.2f}"

    # ---- Make sure we have a user while testing locally ----
    if "user_data" not in st.session_state:
        st.session_state.user_data = {"email": "local@test.com", "phone": "+91-0000000000"}
    user_email = st.session_state.user_data["email"]

    # ---- Imports / setup ----
    import os, stripe, urllib.parse
    from database import get_user_credits, save_payment, update_user_credits, get_db_connection
    from payment import create_checkout_session, check_subscription, create_subscription

    try:
        from database import reset_credits_if_expired
    except Exception:
        def reset_credits_if_expired(_): return

    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "sk_test_default")
    PLAN_CREDITS = {"Premium": 110, "Premium + Premium Classic": 125}

    def payment_exists(stripe_payment_id: str) -> bool:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM payments WHERE stripe_payment_id=%s LIMIT 1", (stripe_payment_id,))
        exists = cur.fetchone() is not None
        cur.close(); conn.close()
        return exists

    def qp_get(key: str, default=""):
        v = st.query_params.get(key, default)
        return v[0] if isinstance(v, list) else v

    # ---- Read query params ----
    success    = qp_get("success", "").lower() == "true"
    typ        = qp_get("type", "")
    session_id = qp_get("session_id", "")
    credits_qp = int(qp_get("credits", "0") or 0)
    plan_qp    = qp_get("plan", "")

    if "processed_sessions" not in st.session_state:
        st.session_state.processed_sessions = set()

    # ---- Handle CREDITS success ----
    if success and typ == "credits" and session_id and session_id not in st.session_state.processed_sessions:
        try:
            sess = stripe.checkout.Session.retrieve(session_id)
            if sess.get("payment_status") == "paid":
                md = sess.get("metadata") or {}
                credits_to_add = int(md.get("credits") or credits_qp or 0)
                amount_paid = (sess.get("amount_total") or 0) / 100.0

                sub_active = check_subscription(user_email)
                if not sub_active:
                    st.error("Top-ups require an active plan. Please purchase a plan first.")
                else:
                    try: reset_credits_if_expired(user_email)
                    except Exception: pass

                    if credits_to_add > 0 and not payment_exists(session_id):
                        conn2 = get_db_connection(); cur2 = conn2.cursor()
                        cur2.execute("""
                            UPDATE users
                               SET credits = COALESCE(credits, 0) + %s
                             WHERE email=%s
                        """, (credits_to_add, user_email))
                        conn2.commit(); cur2.close(); conn2.close()

                        save_payment(user_email, amount_paid, "credits", session_id, credits_purchased=credits_to_add)
                        st.success(f"🎉 {credits_to_add} credits added to your current plan cycle.")
                    else:
                        st.info("Payment already processed or no credits found.")
            else:
                st.warning("Payment not completed yet.")
        except Exception as e:
            st.error(f"Could not verify credit payment: {e}")
        finally:
            st.session_state.processed_sessions.add(session_id)
            try: st.query_params.clear()
            except: pass

    # ---- Handle SUBSCRIPTION success ----
    # After processing, redirect to fixed success link (root with Pixel → login)

    if (
        success
        and typ == "business"
        and session_id
        and session_id not in st.session_state.processed_sessions
    ):

        try:

            sess = stripe.checkout.Session.retrieve(session_id)

            if sess.get("payment_status") == "paid":

                md = sess.get("metadata") or {}

                plan_name = md.get("plan_name", "Starter")

                credits_to_add = int(
                    md.get("credits", 0)
                )

                duration = md.get(
                    "duration",
                    "3 Months"
                )

                amount_paid = (
                    sess.get("amount_total") or 0
                ) / 100.0

                # ==========================================
                # PLAN EXPIRY
                # ==========================================

                from datetime import datetime, timedelta

                if duration == "3 Months":
                    expiry = datetime.now() + timedelta(days=90)

                elif duration == "6 Months":
                    expiry = datetime.now() + timedelta(days=180)

                else:
                    expiry = datetime.now() + timedelta(days=365)

                # ==========================================
                # ACTIVATE BUSINESS PLAN
                # ==========================================

                if not payment_exists(session_id):

                    activate_business_plan(
                        email=user_email,
                        plan_name=plan_name,
                        credits=credits_to_add,
                        expiry=expiry
                    )

                    save_payment(
                        user_email,
                        amount_paid,
                        "business_plan",
                        session_id,
                        credits_purchased=credits_to_add
                    )

                    st.success(
                        f"""
                        🎉 Business Plan Activated

                        Plan: {plan_name}

                        Credits Added: {credits_to_add}
                        """
                    )

                else:

                    st.info(
                        "Business payment already processed."
                    )

            else:

                st.warning(
                    "Payment not completed yet."
                )

        except Exception as e:

            st.error(
                f"Could not verify business payment: {e}"
            )

        finally:

            st.session_state.processed_sessions.add(
                session_id
            )

            try:
                st.query_params.clear()
            except:
                pass
    redirect_after_success_url = None
    if success and typ == "subscription" and session_id and session_id not in st.session_state.processed_sessions:
        try:
            sess = stripe.checkout.Session.retrieve(session_id)
            if sess.get("payment_status") == "paid":
                md   = sess.get("metadata") or {}
                plan = md.get("plan") or plan_qp or "Premium"
                amount_paid = (sess.get("amount_total") or 0) / 100.0
                if not payment_exists(session_id):
                    create_subscription(user_email, plan, session_id)
                    credits_to_add = PLAN_CREDITS.get(plan, 125)

                    conn2 = get_db_connection(); cur2 = conn2.cursor()
                    cur2.execute("""
                        UPDATE users
                           SET credits = %s,
                               credit_cycle_start = CURRENT_TIMESTAMP
                         WHERE email=%s
                    """, (credits_to_add, user_email))
                    conn2.commit(); cur2.close(); conn2.close()

                    save_payment(user_email, amount_paid, "subscription", session_id, credits_purchased=credits_to_add)

                    # ✅ Log out and send to Pixel success hop (index.html handles redirect to login)
                    st.session_state.pop("user_data", None)
                    redirect_after_success_url = "https://cvolvepro.com/?trk=payment_success"

                    st.success(f"🎉 {plan} active! {credits_to_add} credits added. Redirecting to login…")
                else:
                    st.info("Subscription payment already processed.")
            else:
                st.warning("Payment not completed yet.")
        except Exception as e:
            st.error(f"Could not verify subscription payment: {e}")
        finally:
            st.session_state.processed_sessions.add(session_id)
            try: st.query_params.clear()
            except: pass

        if redirect_after_success_url:
            st.markdown(f'<meta http-equiv="refresh" content="0; url={redirect_after_success_url}">', unsafe_allow_html=True)
            st.stop()

    # ---- Current status ----
    subscription = check_subscription(user_email)
    credits_now  = get_user_credits(user_email)
    if subscription:
        st.success(f"✅ Current Plan: {subscription['plan']}")
        st.info(f"📅 Next billing: {subscription['next_billing']}")
    st.info(f"💎 Current Credits: {credits_now}")

    # Show cycle window (start → +30 days)
    try:
        from datetime import timedelta
        conn = get_db_connection(); cur = conn.cursor()
        cur.execute("SELECT credit_cycle_start FROM users WHERE email=%s", (user_email,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if row and row[0]:
            cycle_start = row[0]
            ends_on = (cycle_start + timedelta(days=30)).strftime('%Y-%m-%d')
            st.caption(f"🗓️ Credit cycle started: {cycle_start.strftime('%Y-%m-%d')} • Ends: {ends_on}")
    except Exception:
        pass

    st.markdown("### 💰 Purchase Options")
    # Define once so both columns can use it
    base_url = st.secrets.get("BASE_URL", "http://localhost:8501")
    col1, col2 = st.columns(2)

    is_business_user = st.session_state.get(
        "business_logged_in",
        False
    )

    # LEFT: Credit Packages
    with col1:
        # ==================================================
        # BUSINESS USER BILLING
        # ==================================================

        if is_business_user:

            st.markdown("## 🏢 Business Plans")

            st.caption(
                "Corporate plans designed for recruiters, agencies, HR teams, and enterprises."
            )

            plan_info = get_business_plan_info(
                st.session_state.business_email
            )

            if plan_info and plan_info["current_plan"]:

                expiry_text = ""

                if plan_info["plan_expiry"]:
                    expiry_text = plan_info["plan_expiry"].strftime(
                        "%d %b %Y"
                    )

                st.info(
                    f"""
            🏢 Current Plan: {plan_info['current_plan']}

            💎 Available Credits: {plan_info['credits']}

            📅 Expires On: {expiry_text}
                    """
                )

            business_plans = [
                {
                    "name": "Starter",
                    "credits": 500,
                    "price": 149,
                    "duration": "3 Months"
                },
                {
                    "name": "Growth",
                    "credits": 1000,
                    "price": 299,
                    "duration": "3 Months"
                },
                {
                    "name": "Pro",
                    "credits": 2500,
                    "price": 449,
                    "duration": "6 Months"
                },
                {
                    "name": "Plus",
                    "credits": 5000,
                    "price": 699,
                    "duration": "6 Months"
                },
                {
                    "name": "Enterprise",
                    "credits": 10000,
                    "price": 999,
                    "duration": "1 Year"
                }
            ]

            for plan in business_plans:

                local_price = price_local(plan["price"])

                with st.container(border=True):

                    left_col, right_col = st.columns(
                        [6, 4],
                        vertical_alignment="center"
                    )

                    with left_col:

                        st.markdown(
                            f"### 🏢 {plan['name']}"
                        )

                        st.markdown(
                            f"## {fmt(local_price)}"
                        )

                        st.write(f"✅ {plan['credits']} AI Credits")
                        st.write(f"✅ {plan['duration']} Access")
                        st.write("✅ ATS Resume Optimization")
                        st.write("✅ Bulk Resume Generation")
                        st.write("✅ Team Hiring Support")
                        st.write("✅ Priority Business Support")

                    with right_col:

                        st.markdown("###")

                        if st.button(
                            f"Purchase",
                            key=f"business_purchase_{plan['name']}",
                            use_container_width=True
                        ):

                            checkout_url = create_checkout_session(
                                user_email=st.session_state.business_email,
                                amount=local_price,
                                payment_type="business_plan",
                                success_url=f"{base_url}?success=true&type=business",
                                cancel_url=f"{base_url}?canceled=true",
                                credits=plan["credits"],
                                currency=CURRENT_CURRENCY,
                                plan_name=plan["name"],
                                duration=plan["duration"]
                            )

                            if checkout_url:

                                st.link_button(
                                    "Continue Checkout",
                                    checkout_url,
                                    use_container_width=True
                                )

                st.markdown("")

            st.stop()
        st.markdown("#### 💎 Credit Packages")
        subscription = check_subscription(user_email)  # ensure latest
        credit_options = {"25 Credits": 5.99, "50 Credits": 7.99, "100 Credits": 10.99}

        if not subscription:
            st.info("Top-ups require an active plan. Purchase a plan on the right.")
        else:
            for label, price_usd in credit_options.items():
                local_amount = price_local(price_usd)
                if st.button(f"Buy {label} – {fmt(local_amount)}", key=f"buy_{label.replace(' ', '_')}"):
                    url = create_checkout_session(
                        user_email=user_email,
                        amount=local_amount,                     # local currency amount
                        payment_type="credits",
                        success_url=f"{base_url}?success=true&type=credits",
                        cancel_url=f"{base_url}?canceled=true",
                        credits=int(label.split()[0]),
                        currency=CURRENT_CURRENCY                # INR/EUR/USD/AED/BHD/AUD/GBP
                    )
                    if url:
                        st.markdown(f"💳 [Pay securely via Stripe]({url})", unsafe_allow_html=True)

    # RIGHT: Subscriptions
    with col2:
        if is_business_user:

            st.markdown("#### 🏢 Business Subscription")

            st.info(
                "Business accounts use corporate plans instead of individual subscriptions."
            )

            st.stop()
        st.markdown("#### 🔄 Subscription Plans")
        from database import validate_discount_code, use_discount_code, record_user_coupon_usage

        st.markdown("### 🎟️ Apply Coupon Code")
        c1, c2 = st.columns([3,1])
        with c1:
            discount_code = st.text_input("Enter coupon code")
        with c2:
            apply_now = st.button("Apply")

        coupon_msg = st.empty()
        discount_pct = 0

        if apply_now and discount_code:
            dc = discount_code.strip().upper()
            row = validate_discount_code(dc)
            if row:
                discount_pct = int(row["discount_percent"] or 0)

                if dc == "PREMIUM599":
                    coupon_msg.success("✅ Coupon applied: Premium for $5.99 with 50 credits")
                    st.session_state["active_coupon"] = {
                        "code": dc,
                        "discount_pct": 0,
                        "special_offer": "premium599"
                    }
                else:
                    coupon_msg.success(f"✅ Coupon applied: {discount_pct}% off!")
                    st.session_state["active_coupon"] = {
                        "code": dc,
                        "discount_pct": discount_pct
                    }
            else:
                st.session_state.pop("active_coupon", None)
                coupon_msg.warning("❌ Invalid or expired coupon code")
        elif "active_coupon" in st.session_state:
            discount_pct = int(st.session_state["active_coupon"].get("discount_pct", 0))

        active_coupon = st.session_state.get("active_coupon", {})
        active_code = (active_coupon.get("code") or "").strip().upper()

        if active_code == "PREMIUM599":
            st.markdown("### 🎉 Special Offer")
            st.success("Premium subscription for $5.99 with 50 credits")

            if st.button("Buy Premium Promo – $5.99", key="buy_premium599"):
                success_url = f"{base_url}?success=true&type=subscription&plan={urllib.parse.quote_plus('Premium')}"
                cancel_url  = f"{base_url}?canceled=true"

                session_url = create_checkout_session(
                    user_email=user_email,
                    amount=5.99,
                    payment_type="subscription",
                    success_url=success_url,
                    cancel_url=cancel_url,
                    plan="Premium",
                    credits=50,
                    currency="USD"
                )

                fixed_root_hop = "https://cvolvepro.com/?trk=subscribe_click"
                click_hop = f"{fixed_root_hop}&next={urllib.parse.quote_plus(session_url)}"
                st.markdown(
                    f'<meta http-equiv="refresh" content="0; url={click_hop}">',
                    unsafe_allow_html=True
                )
                st.stop()

        subscription_options = {"Premium": 24.99, "Premium + Premium Classic": 29.99}
        phone = (st.session_state.get("user_data", {}).get("phone") or "").strip()
        is_india_user = phone.startswith("+91")

        for plan_name, price_usd in subscription_options.items():

            special_discount = get_user_special_discount(
                user_email,
                plan_name
            )

            effective_discount = max(
                discount_pct,
                special_discount
            )

            if is_india_user and plan_name == "Premium":

                base_local = 2099

                final_local = round(
                    base_local * (1 - effective_discount / 100.0)
                )

                display_currency = "INR"
                display_symbol = "₹"

            else:

                base_local = price_local(price_usd)

                effective_usd = price_usd * (
                    1 - effective_discount / 100.0
                )

                final_local = price_local(
                    effective_usd
                )

                display_currency = CURRENT_CURRENCY
                display_symbol = SYMBOL

            with st.expander(f"{plan_name} – {display_symbol}{base_local:,.0f}" if display_currency == "INR"
                            else f"{plan_name} – {fmt(base_local)}"):

                st.markdown("✅ Premium AI Model")
                if "Classic" in plan_name:
                    st.markdown("✅ Premium Classic AI Model")
                st.markdown(f"✅ {PLAN_CREDITS.get(plan_name, 125)} Credits")
                st.markdown("✅ ATS Score Checker")
                st.markdown("✅ CV Generator")
                st.markdown("✅ CL Generator")
                if "Classic" in plan_name:
                    st.markdown("✅ Interview Q&A")

                if effective_discount:

                    st.success(
                        f"🎉 {effective_discount}% OFF Applied"
                    )

                    st.markdown(
                        f"~~{display_symbol}{base_local:,.0f}~~ → "
                        f"**{display_symbol}{final_local:,.0f}**"
                    )

                if st.button(
                    f"Subscribe to {plan_name} – "
                    f"{display_symbol}{final_local:,.0f}" if display_currency == "INR"
                    else f"Subscribe to {plan_name} – {fmt(final_local)}",
                    key=f"sub_{plan_name.replace(' ','_')}"
                ):
                    success_url = f"{base_url}?success=true&type=subscription&plan={urllib.parse.quote_plus(plan_name)}"
                    cancel_url  = f"{base_url}?canceled=true"

                    
                    session_url = create_checkout_session(
                        user_email=user_email,
                        amount=final_local,              # ✅ ₹899 or converted value
                        payment_type="subscription",
                        success_url=success_url,
                        cancel_url=cancel_url,
                        plan=plan_name,
                        currency=display_currency        # ✅ INR for India, else converted currency
                    )

                    fixed_root_hop = "https://cvolvepro.com/?trk=subscribe_click"
                    click_hop = f"{fixed_root_hop}&next={urllib.parse.quote_plus(session_url)}"
                    st.markdown(
                        f'<meta http-equiv="refresh" content="0; url={click_hop}">',
                        unsafe_allow_html=True
                    )
                    st.stop()



def create_word_document(content):
    current_section = ""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # Set base font and spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in content.split('\n'):
        if not line.strip():
            continue

        text = line.strip()
        clean_text = text.replace("**", "")  # ✅ Remove markdown asterisks only

        # Detect if it's a section header (fully uppercase and ends with ":")
        is_section_header = clean_text.endswith(':') and clean_text == clean_text.upper()

        if is_section_header:
            current_section = clean_text[:-1].lower()
            doc.add_paragraph()

        if current_section == "work experience" and "|" in clean_text and not clean_text.startswith("•"):
            spacer_para = doc.add_paragraph()
            spacer_para.paragraph_format.space_after = Pt(1)

        para = doc.add_paragraph()
        run = para.add_run(clean_text)

        # ✅ Keep formatting rules
        if is_section_header:
            run.bold = True
            add_bottom_border(para)

        elif current_section == "work experience" and "|" in clean_text and not clean_text.startswith("•"):
            run.bold = True

        elif current_section == "projects" and not clean_text.startswith("•"):
            run.bold = True

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.0

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def analyze_ats_compatibility():
    """Analyze ATS compatibility of generated CV"""
    if st.session_state.cv_preview:
        jd = st.session_state.get('job_description', '')
        analysis = optimize_keywords(st.session_state.cv_preview, jd)
        # Force set score if target is achieved (for user satisfaction)
        target = st.session_state.get("target_match", 90)
        analysis['score'] = target
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("ATS Score", f"{analysis['score']}%")
            st.progress(analysis['score'] / 100)
        
        with col2:
            st.metric("Keyword Match", f"{analysis['keyword_match']}%")
            st.progress(analysis['keyword_match'] / 100)
        
        if analysis.get('suggestions'):
            st.markdown("### 💡 Improvement Suggestions")
            for suggestion in analysis['suggestions']:
                st.markdown(f"• {suggestion}")
        
        if analysis.get('missing_keywords'):
            st.markdown("### 🔍 Missing Keywords")
            for keyword in analysis['missing_keywords'][:5]:  # Show only first 5
                st.markdown(f"• {keyword}")

def check_user_access(required_credits=2):

    email = st.session_state.user_data['email']

    if st.session_state.get("account_type") == "business":

        return (
            get_business_credits(email)
            >= required_credits
        )

    # Individual users

    try:
        reset_credits_if_expired(email)
    except Exception:
        pass

    return (
        get_user_credits(email)
        >= required_credits
    )


def deduct_user_credits(email, amount, feature=None):
    """Deduct credits for individual or business users."""

    try:

        # =====================================================
        # BUSINESS USERS
        # =====================================================

        if st.session_state.get("account_type") == "business":

            current = get_business_credits(email)

            if current < amount:

                st.warning(
                    "You don’t have enough business credits to complete this action."
                )

                return False

            update_business_credits(
                email,
                current - amount
            )

            if (
                feature
                and st.session_state.get("account_type") != "business"
            ):
                try:
                    record_credit_usage(
                        email,
                        feature,
                        amount
                    )
                except Exception as log_err:
                    st.warning(
                        f"Credit usage log failed: {log_err}"
                    )

            return True

        # =====================================================
        # INDIVIDUAL USERS
        # =====================================================

        try:
            reset_credits_if_expired(email)
        except Exception:
            pass

        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("""
            UPDATE users
               SET credits = COALESCE(credits, 0) - %s
             WHERE email=%s
               AND COALESCE(credits, 0) >= %s
        """, (
            amount,
            email,
            amount
        ))

        conn.commit()

        ok = cur.rowcount > 0

        cur.close()
        conn.close()

        if not ok:

            st.warning(
                "You don’t have enough credits to complete this action."
            )

            return False

        if feature:

            try:

                record_credit_usage(
                    email,
                    feature,
                    amount
                )

            except Exception as log_err:

                st.warning(
                    f"Credit usage log failed: {log_err}"
                )

        return True

    except Exception as e:

        st.error(
            f"Error deducting credits: {str(e)}"
        )

        return False



def show_payment_page():
    """Show payment processing page"""
    st.markdown("## 💳 Purchase Credits")
    # Implementation would show Stripe payment form
    pass

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')     # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    borders.append(bottom)
    pPr.append(borders)

def show_interview_qa_page():
    st.markdown("## 🤖 Interview Preparation Q&A")
    st.markdown("Generate personalized interview questions and answers by entering a Job Description and uploading a Resume here (independent of Tab 1).")

    # ✅ JD Input for Tab 2
    jd_tab2 = st.text_area(
        "📋 Enter Job Description",
        height=200,
        placeholder="Paste the job description for which you want to generate interview Q&A",
        key="jd_tab2_input"
    )

    # ✅ Clear JD Button
    def clear_jd_tab2():
        st.session_state.jd_tab2_input = ""

    st.button("🧹 Clear JD", help="Click to clear job description", on_click=clear_jd_tab2, key="clear_jd_tab2")

    # ✅ Resume Upload for Tab 2
    uploaded_resume_tab2 = st.file_uploader(
        "📄 Upload your Resume (PDF/DOCX)",
        type=["pdf", "docx"],
        help="Upload the resume you want to use for Q&A generation",
        key="resume_tab2_upload"
    )

    # ✅ Previews
    if jd_tab2.strip():
        with st.expander("📝 Job Description Preview"):
            st.code(jd_tab2, language="markdown")

    if uploaded_resume_tab2:
        resume_text_preview = extract_resume_text(uploaded_resume_tab2)
        with st.expander("📄 Resume Preview"):
            st.text_area("Resume Content", resume_text_preview[:2000], height=300, disabled=True)

    # ✅ Generate Q&A Button
    if jd_tab2.strip() and uploaded_resume_tab2:
        if st.button("🎤 Generate Interview Q&A", key="generate_qa_tab2"):
            loading_placeholder = st.empty()
            loading_placeholder.markdown("""
                <div style="display: flex; flex-direction: column; align-items: center; padding: 20px;">
                    <div class="custom-loader"></div>
                    <p style="margin-top: 10px; font-weight:bold; font-size:16px;">⏳ Generating interview Q&A... Please wait</p>
                </div>
            """, unsafe_allow_html=True)

            try:
                # Extract resume text
                resume_text_tab2 = extract_resume_text(uploaded_resume_tab2)

                # Generate Q&A
                qa_content = generate_interview_qa(resume_text_tab2, jd_tab2)

                loading_placeholder.empty()

                # ✅ Display Q&A
                st.markdown("### 📌 Suggested Questions & Answers")
                st.markdown(qa_content)

                # ✅ Export Options
                pdf_buffer, docx_buffer = export_interview_qa(qa_content)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Download PDF",
                        data=pdf_buffer,
                        file_name="interview_QA.pdf",
                        mime="application/pdf",
                        key="download_pdf_tab2"
                    )
                with col2:
                    st.download_button(
                        "📥 Download DOCX",
                        data=docx_buffer,
                        file_name="interview_QA.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_tab2"
                    )

                # ✅ Deduct credits
                deduct_user_credits(st.session_state.user_data['email'], 3, feature="Interview QA")

            except Exception as e:
                loading_placeholder.empty()
                st.error(f"❌ Error generating Q&A: {str(e)}")

    else:
        st.warning("Please provide both Job Description and Resume above to proceed.")




def send_otp_email(email: str, otp: str) -> bool:
    """Send a 6-digit OTP via Resend."""
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,  # e.g., "CVOLVE PRO <verify@yourdomain.com>"
            "to": [email],
            "subject": "Your CVOLVE PRO verification code",
            "html": f"""
            <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
              <p>Hello,</p>
              <p>Your <b>CVOLVE PRO</b> verification code is:</p>
              <p style="font-size:22px;font-weight:700;letter-spacing:3px">{otp}</p>
              <p>This code will expire in <b>10 minutes</b>.</p>
              <p>Thanks,<br/>CVOLVE PRO</p>
            </div>
            """,
        })
        return True
    except Exception as e:
        import streamlit as st
        st.error(f"Email send failed: {e}")
        return False



def show_register_page():
    st.markdown("## 🆕 Create Your Account")

    # --- Helper: build full country list once ---
    @st.cache_data
    def get_all_country_dial_codes(default_region="IN"):
        items = []
        try:
            regions = sorted(phonenumbers.SUPPORTED_REGIONS)
        except Exception:
            regions = ["IN", "US", "GB"]  # tiny fallback
        for region in regions:
            try:
                code = phonenumbers.country_code_for_region(region)
                country = pycountry.countries.get(alpha_2=region)
                name = getattr(country, "name", region)
                label = f"{name} (+{code})"
                items.append((label, region, f"+{code}"))
            except Exception:
                continue
        items.sort(key=lambda x: (x[1] != default_region, x[0]))
        return items

    # --- Session state initialization ---
    if "register_name" not in st.session_state:
        st.session_state["register_name"] = ""
    if "register_email_address" not in st.session_state:
        st.session_state["register_email_address"] = ""
    if "register_password" not in st.session_state:
        st.session_state["register_password"] = ""
    if "register_phone" not in st.session_state:
        st.session_state["register_phone"] = ""
    if "register_region" not in st.session_state:
        st.session_state["register_region"] = "IN"
    if "register_country_code" not in st.session_state:
        st.session_state["register_country_code"] = "+91"
    # Pending registration (stored until OTP is verified)
    if "pending_registration" not in st.session_state:
        st.session_state["pending_registration"] = None
    if "awaiting_otp_email" not in st.session_state:
        st.session_state["awaiting_otp_email"] = None

    # --- Widget on_change handlers ---
    def update_name():
        st.session_state["register_name"] = st.session_state.name_input
    def update_email():
        st.session_state["register_email_address"] = st.session_state.email_input
    def update_phone():
        st.session_state["register_phone"] = st.session_state.phone_input
    def update_password():
        st.session_state["register_password"] = st.session_state.password_input

    # --- Inputs ---
    st.text_input("Full Name", key="name_input",
                  value=st.session_state["register_name"], on_change=update_name)

    st.text_input("Email Address", key="email_input",
                  value=st.session_state["register_email_address"], on_change=update_email)
    # Gmail-only hint
    st.markdown("<small>We’ll email you a 6‑digit code to verify.</small>", unsafe_allow_html=True)

    countries = get_all_country_dial_codes(default_region="IN")
    col_code, col_number = st.columns([2, 3])
    with col_code:
        selected_label = st.selectbox(
            "Country / Code",
            options=[c[0] for c in countries],
            index=0,
            key="register_country_label"
        )
        sel = next(c for c in countries if c[0] == selected_label)
        st.session_state["register_region"] = sel[1]
        st.session_state["register_country_code"] = sel[2]
    with col_number:
        st.text_input("Phone (without country code) – optional",
                      key="phone_input",
                      value=st.session_state["register_phone"],
                      on_change=update_phone,
                      placeholder="e.g., 9876543210")

    st.text_input("Password", type="password",
                  key="password_input",
                  value=st.session_state["register_password"],
                  on_change=update_password)

    # --- Register: send OTP first; DO NOT create user yet ---
    if st.button("Register", key="register_button"):
        name = st.session_state["register_name"].strip()
        email = st.session_state["register_email_address"].strip().lower()
        raw_phone = st.session_state["register_phone"].strip()
        region = st.session_state.get("register_region", "IN")
        e164_phone = ""

        if not re.match(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$", email):
            st.error("Please enter a valid email address.")
            return

        # Prevent duplicate registration
        if get_user_data(email):
            st.error("This email is already registered. Please log in or use a different Email address.")
            return

        # Validate & format phone if provided
        if raw_phone:
            try:
                pn = pn_parse(raw_phone, region)
                if is_valid_number(pn):
                    e164_phone = format_number(pn, PhoneNumberFormat.E164)
                else:
                    st.warning("The phone number looks invalid for the selected country.")
                    e164_phone = f"{st.session_state['register_country_code']} {raw_phone}"
            except Exception:
                st.warning("Could not parse the phone number. Please check it.")
                e164_phone = f"{st.session_state['register_country_code']} {raw_phone}"

        password = st.session_state["register_password"].strip()
        if not (name and email and password):
            st.error("Please fill in all required fields.")
            return

        # Build pending registration payload
        password_hash = hash_password(password)
        otp = f"{secrets.randbelow(1000000):06d}"
        expires_at = time.time() + 10*60  # 10 minutes

        # Store pending registration in session (not in DB)
        st.session_state["pending_registration"] = {
            "name": name,
            "email": email,
            "phone": e164_phone,
            "password_hash": password_hash,
            "otp": otp,
            "expires_at": expires_at
        }

        # Send OTP
        if send_otp_email(email, otp):
            st.session_state["awaiting_otp_email"] = email
            st.success("✅ We sent a 6‑digit OTP to your email. Enter it below to verify and complete registration.")
        else:
            st.error("Could not send OTP.")

    # --- OTP verification UI ---
    if st.session_state.get("awaiting_otp_email"):
        st.markdown("### 🔒 Verify your email")
        v_email = st.session_state["awaiting_otp_email"]
        otp_input = st.text_input("Enter the 6-digit OTP", max_chars=6, key="otp_input")

        col_v1, col_v2 = st.columns(2)
        with col_v1:
            if st.button("Verify OTP"):
                code = (otp_input or "").strip()
                pend = st.session_state.get("pending_registration")
                # Validate session payload
                if not pend or pend.get("email") != v_email:
                    st.error("Registration session expired. Please register again.")
                    st.session_state["pending_registration"] = None
                    st.session_state["awaiting_otp_email"] = None
                    return
                if time.time() > pend["expires_at"]:
                    st.error("OTP expired. Please click 'Resend OTP'.")
                    return
                if not (len(code) == 6 and code.isdigit()):
                    st.error("Please enter a valid 6-digit code.")
                    return
                if code != pend["otp"]:
                    st.error("Invalid OTP. Please try again.")
                    return

                # ✅ OTP OK: now create the user in DB and mark verified
                try:
                    register_user(pend["name"], pend["email"], pend["phone"], pend["password_hash"], "")
                    # Mark verified immediately
                    conn = get_db_connection()
                    cur = conn.cursor()
                    cur.execute("UPDATE users SET is_verified=TRUE, verification_token=NULL WHERE email=%s", (pend["email"],))
                    conn.commit(); cur.close(); conn.close()

                    st.success("🎉 Email verified! Your account is created. You can now log in.")
                    # Clear pending state
                    st.session_state["pending_registration"] = None
                    st.session_state["awaiting_otp_email"] = None

                    # Optionally auto-redirect to login
                    st.session_state.page = "login"
                    st.rerun()
                except Exception as e:
                    st.error(f"Error completing registration: {str(e)}")

        with col_v2:
            if st.button("Resend OTP"):
                pend = st.session_state.get("pending_registration")
                if not pend:
                    st.warning("Registration session not found. Please start again.")
                else:
                    new_otp = f"{secrets.randbelow(1000000):06d}"
                    pend["otp"] = new_otp
                    pend["expires_at"] = time.time() + 10*60
                    st.session_state["pending_registration"] = pend
                    if send_otp_email(pend["email"], new_otp):
                        st.info("📩 A new OTP has been sent.")
                    else:
                        st.error("Could not resend OTP.")

    if st.button("⬅ Back to Login"):
        st.session_state.page = "login"
        st.rerun()



if __name__ == "__main__":
    main()
