"""
CVOLVE PRO — AI Interview Practice Module
==========================================
Handles the full interview practice session:
 - Structured Q&A generation (Behavioral / Technical, 3 difficulty levels)
 - AI interviewer (question-by-question flow)
 - Voice answer recording + transcription (browser MediaRecorder via ST component)
 - AI evaluation (semantic, keyword, structure, completeness)
 - Downloadable feedback report (PDF + DOCX)
"""

import streamlit as st
import json
import re
import os
from io import BytesIO
from datetime import datetime

import google.generativeai as genai
import openai
from streamlit import session_state as st_session
from tts_utils import tts_component_html

# ─────────────────────────────────────────────────────────────────────────────
# Gemini model (shared with cv_generator)
# ─────────────────────────────────────────────────────────────────────────────
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
_model = genai.GenerativeModel("gemini-2.5-flash")
openai.api_key = os.getenv("OPENAI_API_KEY")


# ─────────────────────────────────────────────────────────────────────────────
# Credit costs per duration
# ─────────────────────────────────────────────────────────────────────────────
DURATION_CREDITS = {
    "15 minutes": 5,
    "30 minutes": 8,
    "45 minutes": 12,
}

# Fixed: 15 questions total — 5 General, 5 Technical (JD-based), 5 Resume-based
QUESTION_COUNTS = {
    "general": 5,
    "technical": 5,
    "resume": 5,
}


# ─────────────────────────────────────────────────────────────────────────────
# Demo AI fallback (works without any API key)
# ─────────────────────────────────────────────────────────────────────────────
def _demo_ai_response(prompt: str, json_mode: bool) -> str:
    import hashlib
    import json as _json

    seed = hashlib.md5(prompt.encode()).hexdigest()

    if "interview question bank" in prompt or "generate_structured_interview_qa" in prompt.lower():
        # Extract resume and JD from prompt to seed variety
        parts = prompt.split("JOB DESCRIPTION:")
        jd_part = parts[1].split("RESUME:")[0].strip() if len(parts) > 1 else ""
        resume_part = prompt.split("RESUME:")[-1].strip() if "RESUME:" in prompt else ""
        seed = hashlib.md5((jd_part + resume_part).encode()).hexdigest()
        # Use seed to pick from question pools
        rng = int(seed[:8], 16)

        general_pool = [
            {"question": "Tell me about yourself and why you're interested in this role.", "ideal_answer": "I have extensive experience in areas relevant to this role. My background includes key achievements that demonstrate my capability. I am excited about this opportunity because it aligns with my skills and career goals.", "key_points": ["Relevant experience", "Skills match", "Enthusiasm for role", "Career narrative", "Specific examples"]},
            {"question": "Describe a time you worked successfully as part of a team.", "ideal_answer": "In my previous role, I collaborated with cross-functional teams on a complex project. I contributed by taking ownership of key deliverables, which led to measurable positive outcomes for the business.", "key_points": ["Team collaboration", "Specific role", "Measurable outcome", "Communication", "Conflict resolution"]},
            {"question": "Tell me about a time you had to manage competing priorities.", "ideal_answer": "I was managing multiple projects with tight deadlines. I prioritized based on business impact, communicated clearly with stakeholders, and delivered all projects on time by effectively delegating tasks.", "key_points": ["Priority management", "Stakeholder communication", "Delegation", "Time management", "Results"]},
            {"question": "Describe a situation where you influenced a decision without authority.", "ideal_answer": "I identified an opportunity to improve our team's workflow. I built a business case, presented it to leadership with data-driven evidence, and gained buy-in to implement the change, resulting in 20% efficiency gain.", "key_points": ["Influence without authority", "Data-driven approach", "Leadership", "Business case", "Measurable impact"]},
            {"question": "Tell me about a time you failed and what you learned from it.", "ideal_answer": "I launched a feature that didn't gain user adoption. I took ownership, analyzed feedback, and discovered the UX was confusing. I redesigned it based on user research, and adoption increased significantly.", "key_points": ["Ownership of failure", "Analysis", "Learning applied", "Humility", "Improved outcome"]},
            {"question": "Describe a time you had to adapt to a significant change.", "ideal_answer": "Our team underwent a major reorganization. I embraced the change by proactively learning new systems, building relationships with new team members, and maintaining productivity throughout the transition.", "key_points": ["Adaptability", "Proactive learning", "Resilience", "Team collaboration", "Productivity"]},
            {"question": "Tell me about a time you went above and beyond for a project.", "ideal_answer": "I noticed a critical gap in our project plan that could cause delays. I voluntarily took on additional work, worked extra hours, and coordinated with multiple teams to ensure we met the deadline with high quality.", "key_points": ["Initiative", "Dedication", "Cross-team coordination", "Problem-solving", "Quality focus"]},
            {"question": "Describe a time you had to deliver difficult feedback.", "ideal_answer": "A team member's work quality was declining. I scheduled a private meeting, used specific examples to illustrate concerns, and worked with them to create an improvement plan. Performance improved significantly within a month.", "key_points": ["Constructive feedback", "Empathy", "Specific examples", "Improvement plan", "Leadership"]},
        ]
        tech_pool = [
            {"question": "How would you design a scalable data processing pipeline for real-time analytics?", "ideal_answer": "I would use a streaming architecture with Kafka for ingestion, Apache Flink or Spark Streaming for processing, and a time-series database for storage. The system should handle backpressure, support exactly-once semantics, and scale horizontally.", "key_points": ["Streaming architecture", "Message queue", "Stream processing", "Scalability", "Fault tolerance"]},
            {"question": "Explain the trade-offs between REST and GraphQL API design.", "ideal_answer": "REST is simpler with clear resource-oriented endpoints and HTTP caching. GraphQL offers flexible queries but adds complexity in query parsing, authorization per field, and caching. Choose REST for simple CRUD, GraphQL for complex data requirements.", "key_points": ["REST principles", "GraphQL flexibility", "Caching differences", "Complexity trade-offs", "Use case selection"]},
            {"question": "How would you optimize a system with high read and write throughput?", "ideal_answer": "Use CQRS pattern to separate read and write models. Implement read replicas, caching layer with Redis, database sharding for writes, and asynchronous processing for write-heavy operations. Monitor and tune based on actual usage patterns.", "key_points": ["CQRS pattern", "Read replicas", "Caching strategy", "Sharding", "Async processing"]},
            {"question": "Design a system to handle millions of concurrent users for a real-time application.", "ideal_answer": "Use microservices architecture with auto-scaling, load balancers, CDN for static content, WebSocket connections for real-time features, distributed caching, and database sharding. Implement circuit breakers and rate limiting.", "key_points": ["Microservices", "Auto-scaling", "Load balancing", "CDN", "WebSocket", "Caching"]},
            {"question": "Explain how you would implement a recommendation system.", "ideal_answer": "Use collaborative filtering (user-based or item-based) and content-based filtering as complementary approaches. Implement matrix factorization for latent features. Use A/B testing to evaluate performance. Scale with approximate nearest neighbor search.", "key_points": ["Collaborative filtering", "Content-based filtering", "Matrix factorization", "A/B testing", "ANN search"]},
            {"question": "How would you ensure data consistency in a distributed system?", "ideal_answer": "Use distributed transactions with Saga pattern for long-running transactions. Implement eventual consistency where appropriate, with conflict resolution strategies. Use consensus algorithms like Raft for critical operations.", "key_points": ["Saga pattern", "Eventual consistency", "Conflict resolution", "Consensus algorithms", "Transaction management"]},
            {"question": "Design a monitoring and alerting system for a cloud-native application.", "ideal_answer": "Use Prometheus for metrics collection, Grafana for visualization, and Alertmanager for alerting. Implement distributed tracing with Jaeger. Set up structured logging with ELK stack. Define SLOs and alert on error budgets.", "key_points": ["Metrics collection", "Visualization", "Alerting", "Distributed tracing", "Structured logging", "SLOs"]},
            {"question": "Explain approaches for secure API authentication in a microservices architecture.", "ideal_answer": "Use OAuth 2.0 with JWT tokens for stateless authentication. Implement API gateway for centralized auth, short-lived tokens, refresh token rotation, and scoped access. Use mTLS for service-to-service communication.", "key_points": ["OAuth 2.0", "JWT", "API gateway", "Token management", "mTLS", "Scoped access"]},
        ]
        resume_pool = [
            {"question": "Walk me through your most recent role and your key contributions.", "ideal_answer": "In my most recent role, I was responsible for key deliverables. My top achievement involved solving a complex problem using my technical skills, which resulted in measurable business impact.", "key_points": ["Role description", "Key achievements", "Metrics", "Technologies used", "Impact"]},
            {"question": "What specific technologies are you most proficient in and why?", "ideal_answer": "I am most proficient in the technologies I have used extensively in previous projects. I have deep practical experience applying them to solve real-world problems and deliver results.", "key_points": ["Technology proficiency", "Project examples", "Depth of experience", "Practical application", "Results"]},
            {"question": "Describe a challenging problem you solved in a previous project.", "ideal_answer": "On a previous project, we faced a significant technical challenge. I analyzed the problem, designed a solution, implemented it, and the outcome was highly positive for the business.", "key_points": ["Problem identification", "Solution approach", "Technical depth", "Persistence", "Result"]},
            {"question": "How has your previous experience prepared you for this role?", "ideal_answer": "My experience has given me a strong foundation in relevant skills. I have developed expertise through hands-on work and am well-prepared to contribute immediately to this role.", "key_points": ["Experience connection", "Skill transferability", "Career progression", "Self-awareness", "Role alignment"]},
            {"question": "What areas of your expertise would you like to develop further?", "ideal_answer": "I have strong foundational skills but am looking to deepen my expertise in emerging areas relevant to this role. This position offers the opportunity to grow in these areas.", "key_points": ["Honest self-assessment", "Growth mindset", "Role relevance", "Learning goals", "Ambition"]},
            {"question": "Tell me about a project where you had to learn a new technology quickly.", "ideal_answer": "I was assigned to a project requiring a technology I hadn't used before. I created a structured learning plan, built a prototype, and delivered the project successfully within the deadline.", "key_points": ["Quick learning", "Structured approach", "Prototyping", "Delivery", "Adaptability"]},
            {"question": "Describe a time you improved an existing process or system.", "ideal_answer": "I identified inefficiencies in an existing workflow. I proposed and implemented improvements that reduced time by a significant percentage and improved quality.", "key_points": ["Process improvement", "Initiative", "Quantified impact", "Implementation", "Quality improvement"]},
            {"question": "How do you stay current with industry trends and technologies?", "ideal_answer": "I regularly read technical blogs, participate in online courses, attend conferences, and work on side projects. I recently explored new technologies that are directly relevant to this role.", "key_points": ["Continuous learning", "Industry awareness", "Practical application", "Learning methods", "Relevance"]},
        ]

        def pick(questions, count, offset):
            return [questions[(offset + i) % len(questions)] for i in range(count)]

        offset = rng
        demo = {
            "general": pick(general_pool, 5, offset),
            "technical": pick(tech_pool, 5, offset + 5),
            "resume": pick(resume_pool, 5, offset + 10),
        }
        return _json.dumps(demo)

    if "evaluate" in prompt.lower() and ("answer" in prompt.lower() or "question" in prompt.lower()):
        return _json.dumps({
            "score": 28, "meaning_match": 25, "keyword_coverage": 20,
            "keywords_covered": [],
            "keywords_missed": ["all key points"],
            "structure_score": 15, "completeness_score": 10,
            "confidence_indicators": [],
            "strengths": [],
            "improvements": ["Answer was too brief or generic", "Must address specific key points from the ideal answer"],
            "improved_answer": "A strong answer would include specific examples, cover the key points listed, and follow a clear structure.",
            "brief_feedback": "Your answer needs significant improvement. It lacks depth and does not address the key points expected. Review the ideal answer and try again with more detail and specific examples."
        })

    if "feedback" in prompt.lower() or "performance report" in prompt.lower() or "career coach" in prompt.lower():
        return _json.dumps({
            "overall_summary": "Your interview performance needs improvement. Focus on providing structured, detailed answers with specific examples.",
            "key_strengths": ["Attempted to answer questions"],
            "weak_areas": ["Answer depth and specificity", "Coverage of key points", "Structure and clarity"],
            "general_feedback": "Your general answers need more structure. Use the STAR method: describe the Situation, Task, Action, and Result with specific details.",
            "technical_feedback": "Technical answers should demonstrate depth. Explain not just what but how and why. Include specific technologies and approaches.",
            "resume_feedback": "Your resume-based answers must draw specific examples from your actual experience. Include metrics and concrete contributions.",
            "recommendations": [
                "Practice the STAR method with stories from your experience",
                "Study the job description requirements in depth",
                "Record yourself answering questions to identify gaps",
                "Prepare specific metrics for each achievement on your resume"
            ],
            "next_steps": "Review the questions where you scored lowest, study the ideal answers, and retry. Schedule another session after preparation."
        })

    return '{"result": "Demo response"}'


# ─────────────────────────────────────────────────────────────────────────────
# AI call helper (Gemini / OpenAI)
# ─────────────────────────────────────────────────────────────────────────────
_DEMO_MODE = False

def _check_demo_mode():
    global _DEMO_MODE
    if _DEMO_MODE:
        return True
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or ""
    openai_key = os.getenv("OPENAI_API_KEY") or ""
    if not gemini_key and not openai_key:
        _DEMO_MODE = True
        return True
    return False


def _ai_call(prompt: str, json_mode: bool = False) -> str:
    """Route to OpenAI or Gemini based on session setting, with demo fallback."""
    if _check_demo_mode():
        return _demo_ai_response(prompt, json_mode)

    try:
        if st_session.get("ai_model") == "openai":
            resp = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                response_format={"type": "json_object"} if json_mode else {"type": "text"},
            )
            return resp.choices[0].message.content
        else:
            resp = _model.generate_content(
                prompt,
                generation_config={
                    "temperature": 0.7,
                    "response_mime_type": "application/json" if json_mode else "text/plain",
                },
            )
            return resp.text
    except Exception:
        _DEMO_MODE = True
        return _demo_ai_response(prompt, json_mode)


# ─────────────────────────────────────────────────────────────────────────────
# 1. Generate structured interview Q&A
# ─────────────────────────────────────────────────────────────────────────────
def generate_structured_interview_qa(resume_text: str, job_description: str, duration: str) -> dict:
    """
    Returns a dict:
    {
        "general":    [{"question": ..., "ideal_answer": ..., "key_points": [...]}, ...],  (5 questions)
        "technical":  [{"question": ..., "ideal_answer": ..., "key_points": [...]}, ...],  (5 questions)
        "resume":     [{"question": ..., "ideal_answer": ..., "key_points": [...]}, ...]   (5 questions)
    }
    Total: 15 questions across 3 categories.
    """
    gen_n = QUESTION_COUNTS["general"]
    tech_n = QUESTION_COUNTS["technical"]
    res_n = QUESTION_COUNTS["resume"]

    prompt = f"""
You are an expert technical interviewer and career coach. Your task is to generate a highly targeted interview question bank that is DEEPLY ALIGNED with the specific job description below.

OUTPUT FORMAT (strict JSON only, no markdown, no extra text):
{{
  "general": [
    {{
      "question": "...",
      "ideal_answer": "A complete STAR-format answer (200-250 words)...",
      "key_points": ["key point 1", "key point 2", "key point 3", "key point 4", "key point 5"]
    }}
  ],
  "technical": [...],
  "resume": [...]
}}

CRITICAL RULES — Follow these EXACTLY:

**General ({gen_n} questions):**
- These must be behavioral/situational questions that are SPECIFIC to this job's context, NOT generic.
- Example: If the JD mentions "cross-functional collaboration with product managers", ask about a time they navigated disagreement with a product manager on technical requirements.
- Each general question must DIRECTLY reference a specific responsibility, challenge, or context from the job description.

**Technical ({tech_n} questions):**
- EVERY question must test knowledge of a specific tool, language, framework, or concept EXPLICITLY mentioned in the job description.
- Extract ALL technical keywords from the JD first. Then write one question per keyword or combination.
- Example: If JD says "Python, SQL, machine learning, AWS", ask about designing a Python-based ML pipeline on AWS with SQL data sources.
- These must be DEEP technical questions, not surface-level definitions. Probe architecture, trade-offs, design decisions.
- NO generic technical questions. Every question must be IMPOSSIBLE to ask without the specific JD.

**Resume ({res_n} questions):**
- Each question must ask about a SPECIFIC project, role, technology, or achievement mentioned in the resume.
- Probe depth: ask WHY they made certain choices, WHAT the outcome was, HOW they overcame challenges.
- Questions must tie the candidate's past experience to the job requirements.

ADDITIONAL REQUIREMENTS:
- ALL 15 questions must be unique and non-repetitive.
- ideal_answer: 200-250 words, comprehensive, well-structured (STAR for behavioral, technical depth for technical).
- key_points: 5-7 essential concepts/keywords the answer must cover.
- Do NOT include numbering inside question text.
- If the JD or resume is thin, infer reasonable questions based on what IS provided — never fill with generic questions.

JOB DESCRIPTION (must anchor EVERY question):
{job_description}

RESUME (for resume-specific questions):
{resume_text}
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        data = json.loads(raw_clean)
        return data
    except Exception as e:
        raise ValueError(f"Failed to parse AI response as JSON: {e}\n\nRaw: {raw[:500]}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. Flatten Q list for session (ordered: behavioral then technical)
# ─────────────────────────────────────────────────────────────────────────────
def flatten_questions(qa_bank: dict) -> list:
    """Returns flat list of question objects with section metadata."""
    flat = []
    for section in ["general", "technical", "resume"]:
        qs = qa_bank.get(section, [])
        for q in qs:
            flat.append({
                "section": section,
                "difficulty": "",
                "question": q.get("question", ""),
                "ideal_answer": q.get("ideal_answer", ""),
                "key_points": q.get("key_points", []),
            })
    return flat


# ─────────────────────────────────────────────────────────────────────────────
# 3. AI Answer Evaluation
# ─────────────────────────────────────────────────────────────────────────────
def evaluate_answer(question: str, ideal_answer: str, key_points: list,
                    user_answer: str, section: str, difficulty: str) -> dict:
    """
    Returns evaluation dict:
    {
      "score": 0-100,
      "meaning_match": 0-100,
      "keyword_coverage": 0-100,
      "keywords_covered": [...],
      "keywords_missed": [...],
      "structure_score": 0-100,
      "completeness_score": 0-100,
      "confidence_indicators": [...],
      "strengths": [...],
      "improvements": [...],
      "improved_answer": "...",
      "brief_feedback": "..."
    }
    """
    if not user_answer or len(user_answer.strip()) < 10:
        return {
            "score": 0, "meaning_match": 0, "keyword_coverage": 0,
            "keywords_covered": [], "keywords_missed": key_points,
            "structure_score": 0, "completeness_score": 0,
            "confidence_indicators": [],
            "strengths": [], "improvements": ["No answer was provided."],
            "improved_answer": ideal_answer,
            "brief_feedback": "No answer provided. Please attempt to answer all questions."
        }

    prompt = f"""
You are a STRICT interview coach evaluating a candidate's answer. You do NOT inflate scores. You compare what the candidate said against the ideal answer and key points.

QUESTION: {question}
SECTION: {section}

IDEAL ANSWER (reference for comparison):
{ideal_answer}

KEY POINTS the answer should cover:
{json.dumps(key_points)}

CANDIDATE'S ANSWER:
{user_answer}

First, carefully read the question and the ideal answer. Then compare the candidate's answer to the ideal answer. Score based on how closely the candidate's answer matches the meaning and covers the key points.

Output ONLY valid JSON (no markdown):
{{
  "score": <overall 0-100>,
  "meaning_match": <0-100, how closely the meaning matches the ideal>,
  "keyword_coverage": <0-100, % of key points explicitly addressed>,
  "keywords_covered": ["<key point explicitly covered>", ...],
  "keywords_missed": ["<key point NOT addressed>", ...],
  "structure_score": <0-100, logical flow, STAR or clear structure>,
  "completeness_score": <0-100, how complete and detailed>,
  "confidence_indicators": ["<positive indicator>", ...],
  "strengths": ["<strength 1>", "<strength 2>", ...],
  "improvements": ["<specific improvement>", ...],
  "improved_answer": "<a better version of the candidate's actual answer, 150-200 words>",
  "brief_feedback": "<2-3 sentence coach feedback>"
}}

STRICT SCORING RUBRIC:
- 0-20: Answer is blank, gibberish, completely unrelated, or random text with no meaningful connection to the question
- 21-40: Answer touches the general topic but misses most key points; vague or generic
- 41-60: Answer addresses some key points but is incomplete; partial understanding shown
- 61-80: Answer covers most key points with reasonable depth; minor gaps remain
- 81-90: Answer covers nearly all key points with good detail and structure
- 91-100: Answer matches the ideal answer closely; all key points covered with excellent depth and examples

CRITICAL RULES:
- If the candidate wrote random text, nonsense, or copied the question back, score 0-15.
- Do NOT give partial credit for unrelated content. The answer must actually address the key points.
- A vague 1-sentence answer that mentions one key point should score no more than 25.
- "Improved_answer" must start from the candidate's actual answer (if any substance), not from the ideal answer.
- Be STRICT. A score of 50 means the answer was mediocre.
"""

    raw = _ai_call(prompt, json_mode=True)
    try:
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        return json.loads(raw_clean)
    except Exception:
        return {
            "score": 50, "meaning_match": 50, "keyword_coverage": 50,
            "keywords_covered": [], "keywords_missed": [],
            "structure_score": 50, "completeness_score": 50,
            "confidence_indicators": [],
            "strengths": ["Answer provided"], "improvements": ["Could not parse evaluation"],
            "improved_answer": ideal_answer,
            "brief_feedback": "Evaluation could not be parsed. Please try again."
        }


# ─────────────────────────────────────────────────────────────────────────────
# 4. Generate Full Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def generate_feedback_report(session_results: list, duration: str) -> dict:
    """
    session_results: list of {question_obj, evaluation}
    Returns a structured feedback report dict.
    """
    total_score = 0
    general_scores = []
    technical_scores = []
    resume_scores = []
    all_keywords_covered = set()
    all_keywords_missed = set()
    well_answered = []
    incomplete_answers = []

    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        score = ev.get("score", 0)
        total_score += score

        section = q_obj.get("section", "")
        if section == "general":
            general_scores.append(score)
        elif section == "technical":
            technical_scores.append(score)
        else:
            resume_scores.append(score)

        all_keywords_covered.update(ev.get("keywords_covered", []))
        all_keywords_missed.update(ev.get("keywords_missed", []))

        q_text = q_obj.get("question", "")
        if score >= 70:
            well_answered.append({"question": q_text, "score": score})
        else:
            incomplete_answers.append({
                "question": q_text,
                "score": score,
                "improvements": ev.get("improvements", []),
                "improved_answer": ev.get("improved_answer", ""),
            })

    n = len(session_results)
    overall = round(total_score / n) if n > 0 else 0
    general_avg = round(sum(general_scores) / len(general_scores)) if general_scores else 0
    technical_avg = round(sum(technical_scores) / len(technical_scores)) if technical_scores else 0
    resume_avg = round(sum(resume_scores) / len(resume_scores)) if resume_scores else 0

    # Remove keywords from missed if they were covered in other questions
    all_keywords_missed -= all_keywords_covered

    # Performance band
    if overall >= 85:
        band = "Excellent"
        band_color = "🟢"
    elif overall >= 70:
        band = "Good"
        band_color = "🔵"
    elif overall >= 55:
        band = "Average"
        band_color = "🟡"
    else:
        band = "Needs Improvement"
        band_color = "🔴"

    # Generate AI-powered insights
    results_summary = []
    for item in session_results:
        ev = item.get("evaluation", {})
        q_obj = item.get("question_obj", {})
        results_summary.append({
            "question": q_obj.get("question", "")[:100],
            "section": q_obj.get("section", ""),
            "difficulty": q_obj.get("difficulty", ""),
            "score": ev.get("score", 0),
            "strengths": ev.get("strengths", []),
            "improvements": ev.get("improvements", []),
        })

    ai_prompt = f"""
You are a senior career coach reviewing a {duration} mock interview session.

OVERALL SCORE: {overall}/100 ({band})
GENERAL AVG: {general_avg}/100
TECHNICAL AVG: {technical_avg}/100
RESUME-BASED AVG: {resume_avg}/100

SESSION RESULTS SUMMARY:
{json.dumps(results_summary, indent=2)}

Generate a concise, honest, and constructive performance report. Output ONLY valid JSON:
{{
  "overall_summary": "<2-3 sentence overall assessment>",
  "key_strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "weak_areas": ["<weak area 1>", "<weak area 2>"],
  "general_feedback": "<specific feedback on general interview answers>",
  "technical_feedback": "<specific feedback on technical answers>",
  "resume_feedback": "<specific feedback on resume-based answers>",
  "recommendations": ["<actionable recommendation 1>", "<recommendation 2>", "<recommendation 3>", "<recommendation 4>"],
  "next_steps": "<what to focus on in the next practice session>"
}}
"""

    try:
        raw = _ai_call(ai_prompt, json_mode=True)
        raw_clean = re.sub(r"^```[a-z]*\n?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        ai_insights = json.loads(raw_clean)
    except Exception:
        ai_insights = {
            "overall_summary": f"You scored {overall}/100 overall. Review the individual question feedback to improve.",
            "key_strengths": ["Completed the practice session"],
            "weak_areas": ["Review questions scored below 70"],
            "general_feedback": "Work on structuring general answers using the STAR method.",
            "technical_feedback": "Strengthen technical depth with concrete examples.",
            "resume_feedback": "Practice articulating your past experience more clearly.",
            "recommendations": [
                "Practice daily with different job descriptions",
                "Record yourself answering questions",
                "Research the company's tech stack thoroughly",
                "Prepare 3-4 strong STAR stories"
            ],
            "next_steps": "Focus on the questions where you scored below 70 and retry with improved answers."
        }

    return {
        "overall_score": overall,
        "performance_band": band,
        "band_color": band_color,
        "general_score": general_avg,
        "technical_score": technical_avg,
        "resume_score": resume_avg,
        "total_questions": n,
        "duration": duration,
        "well_answered": well_answered,
        "incomplete_answers": incomplete_answers,
        "keywords_covered": sorted(all_keywords_covered),
        "keywords_missed": sorted(all_keywords_missed),
        "session_results": session_results,
        **ai_insights,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 5. Export report to PDF & DOCX
# ─────────────────────────────────────────────────────────────────────────────
def export_feedback_report(report: dict):
    """Returns (pdf_buffer, docx_buffer)."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from docx import Document
    from docx.shared import Pt, RGBColor

    # ─── PDF ────────────────────────────────────────────────────────────────
    pdf_buf = BytesIO()
    doc = SimpleDocTemplate(pdf_buf, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("TitleS", fontSize=22, leading=28, fontName="Helvetica-Bold",
                              textColor=colors.HexColor("#1a1a2e"), spaceAfter=6)
    h2_s = ParagraphStyle("H2S", fontSize=14, leading=18, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#16213e"), spaceAfter=4, spaceBefore=12)
    h3_s = ParagraphStyle("H3S", fontSize=11, leading=14, fontName="Helvetica-Bold",
                           textColor=colors.HexColor("#0f3460"), spaceAfter=3, spaceBefore=8)
    body_s = ParagraphStyle("BodyS", fontSize=10, leading=14, spaceAfter=4)
    bullet_s = ParagraphStyle("BulletS", fontSize=10, leading=14, leftIndent=16, spaceAfter=3)
    score_s = ParagraphStyle("ScoreS", fontSize=28, fontName="Helvetica-Bold",
                             textColor=colors.HexColor("#e94560"), spaceAfter=4)
    sub_s = ParagraphStyle("SubS", fontSize=10, leading=13, textColor=colors.grey, spaceAfter=8)

    story = []

    # Header
    story.append(Paragraph("CVOLVE PRO — Interview Feedback Report", title_s))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}", sub_s))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#e94560")))

    # Score summary table
    story.append(Spacer(1, 10))
    score_data = [
        ["Overall", "General", "Technical", "Resume", "Band"],
        [f"{report['overall_score']}/100", f"{report['general_score']}/100",
         f"{report['technical_score']}/100", f"{report['resume_score']}/100",
         report['performance_band']],
    ]
    score_table = Table(score_data, colWidths=[1.2*inch, 1.1*inch, 1.1*inch, 1.1*inch, 1.1*inch])
    score_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16213e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 11),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f4ff")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#ccddff")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(score_table)
    story.append(Spacer(1, 14))

    # Overall summary
    story.append(Paragraph("📋 Overall Assessment", h2_s))
    story.append(Paragraph(report.get("overall_summary", ""), body_s))

    # Strengths
    story.append(Paragraph("💪 Key Strengths", h2_s))
    for s in report.get("key_strengths", []):
        story.append(Paragraph(f"• {s}", bullet_s))

    # Weak areas
    story.append(Paragraph("⚠️ Areas to Improve", h2_s))
    for w in report.get("weak_areas", []):
        story.append(Paragraph(f"• {w}", bullet_s))

    # Section feedback
    story.append(Paragraph("📋 General Questions Feedback", h2_s))
    story.append(Paragraph(report.get("general_feedback", ""), body_s))
    story.append(Paragraph("⚙️ Technical Questions Feedback", h2_s))
    story.append(Paragraph(report.get("technical_feedback", ""), body_s))
    story.append(Paragraph("📄 Resume-based Questions Feedback", h2_s))
    story.append(Paragraph(report.get("resume_feedback", ""), body_s))

    # Questions answered well
    if report.get("well_answered"):
        story.append(Paragraph("✅ Questions Answered Well", h2_s))
        for item in report["well_answered"]:
            story.append(Paragraph(f"• {item['question']} — Score: {item['score']}/100", bullet_s))

    # Incomplete answers with improved versions
    if report.get("incomplete_answers"):
        story.append(Paragraph("📌 Questions Needing Improvement", h2_s))
        for item in report["incomplete_answers"]:
            story.append(Paragraph(item["question"], h3_s))
            story.append(Paragraph(f"Score: {item['score']}/100", sub_s))
            for imp in item.get("improvements", []):
                story.append(Paragraph(f"• {imp}", bullet_s))
            if item.get("improved_answer"):
                story.append(Paragraph("Suggested Improved Answer:", h3_s))
                story.append(Paragraph(item["improved_answer"], body_s))

    # Keywords
    story.append(Paragraph("🔑 Keywords Covered", h2_s))
    covered_text = ", ".join(report.get("keywords_covered", [])) or "None recorded"
    story.append(Paragraph(covered_text, body_s))

    story.append(Paragraph("❌ Keywords Missed", h2_s))
    missed_text = ", ".join(report.get("keywords_missed", [])) or "None — great coverage!"
    story.append(Paragraph(missed_text, body_s))

    # Recommendations
    story.append(Paragraph("🎯 Recommendations for Further Practice", h2_s))
    for r in report.get("recommendations", []):
        story.append(Paragraph(f"• {r}", bullet_s))

    story.append(Paragraph("🚀 Next Steps", h2_s))
    story.append(Paragraph(report.get("next_steps", ""), body_s))

    doc.build(story)
    pdf_buf.seek(0)

    # ─── DOCX ───────────────────────────────────────────────────────────────
    docx_buf = BytesIO()
    wd = Document()
    wd.core_properties.title = "CVOLVE PRO Interview Feedback Report"

    def add_h(doc, text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)

    def add_p(doc, text, bold=False, size=11):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    def add_b(doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    add_h(wd, "CVOLVE PRO — Interview Feedback Report", 1, (26, 26, 46))
    add_p(wd, f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')} | Duration: {report.get('duration', '')}")
    add_h(wd, f"Overall Score: {report['overall_score']}/100 — {report['performance_band']}", 2, (233, 69, 96))
    add_p(wd, f"General: {report['general_score']}/100 | Technical: {report['technical_score']}/100 | Resume: {report['resume_score']}/100 | Questions: {report['total_questions']}")

    add_h(wd, "Overall Assessment", 2)
    add_p(wd, report.get("overall_summary", ""))

    add_h(wd, "Key Strengths", 2)
    for s in report.get("key_strengths", []):
        add_b(wd, s)

    add_h(wd, "Areas to Improve", 2)
    for w in report.get("weak_areas", []):
        add_b(wd, w)

    add_h(wd, "General Questions Feedback", 2)
    add_p(wd, report.get("general_feedback", ""))

    add_h(wd, "Technical Questions Feedback", 2)
    add_p(wd, report.get("technical_feedback", ""))

    add_h(wd, "Resume-based Questions Feedback", 2)
    add_p(wd, report.get("resume_feedback", ""))

    if report.get("well_answered"):
        add_h(wd, "Questions Answered Well", 2)
        for item in report["well_answered"]:
            add_b(wd, f"{item['question']} — {item['score']}/100")

    if report.get("incomplete_answers"):
        add_h(wd, "Questions Needing Improvement", 2)
        for item in report["incomplete_answers"]:
            add_h(wd, item["question"], 3)
            add_p(wd, f"Score: {item['score']}/100")
            for imp in item.get("improvements", []):
                add_b(wd, imp)
            if item.get("improved_answer"):
                add_p(wd, "Suggested Improved Answer:", bold=True)
                add_p(wd, item["improved_answer"])

    add_h(wd, "Keywords Covered", 2)
    add_p(wd, ", ".join(report.get("keywords_covered", [])) or "None")

    add_h(wd, "Keywords Missed", 2)
    add_p(wd, ", ".join(report.get("keywords_missed", [])) or "None — great coverage!")

    add_h(wd, "Recommendations", 2)
    for r in report.get("recommendations", []):
        add_b(wd, r)

    add_h(wd, "Next Steps", 2)
    add_p(wd, report.get("next_steps", ""))

    wd.save(docx_buf)
    docx_buf.seek(0)

    return pdf_buf, docx_buf


# ─────────────────────────────────────────────────────────────────────────────
# 6. Speech-to-text via inline JavaScript (direct textarea fill)
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# 7. Main Streamlit UI — show_interview_practice_page()
# ─────────────────────────────────────────────────────────────────────────────
def show_interview_practice_page(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn):
    """
    Main entry point called from app.py.
    Passes down helper functions to avoid circular imports.
    """
    _init_session()

    # ── Phase router ──────────────────────────────────────────────────────────
    phase = st.session_state.get("interview_phase", "setup")

    if phase == "setup":
        _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn)
    elif phase == "session":
        _phase_session()
    elif phase == "report":
        _phase_report()


def _init_session():
    defaults = {
        "interview_phase": "setup",
        "interview_qa_bank": None,
        "interview_questions_flat": None,
        "interview_current_idx": 0,
        "interview_session_results": [],
        "interview_report": None,
        "interview_duration": "30 minutes",
        "interview_jd": "",
        "interview_resume_text": "",
        "voice_transcript_buffer": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────────────────────────────────────
# Phase 1: Setup
# ─────────────────────────────────────────────────────────────────────────────
def _phase_setup(check_access_fn, deduct_credits_fn, extract_resume_fn, export_qa_fn, generate_qa_fn):

    st.markdown("""
    <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
                border-radius: 16px; padding: 30px; margin-bottom: 24px;">
        <h1 style="color:#e94560; margin:0; font-size:32px;">🤖 AI Interview Practice</h1>
        <p style="color:#a0b4d6; margin:8px 0 0; font-size:16px;">
            A real-time interview simulation powered by AI — generate questions, practice answers,
            and receive a detailed performance report.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # ─── Credit Info Banner ───────────────────────────────────────────────────
    col_c1, col_c2, col_c3 = st.columns(3)
    with col_c1:
        st.markdown("""<div style="background:#fff3cd;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#856404">5 Credits</div>
            <div style="color:#856404;font-size:13px">15-minute session</div></div>""", unsafe_allow_html=True)
    with col_c2:
        st.markdown("""<div style="background:#d1ecf1;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#0c5460">8 Credits</div>
            <div style="color:#0c5460;font-size:13px">30-minute session</div></div>""", unsafe_allow_html=True)
    with col_c3:
        st.markdown("""<div style="background:#d4edda;border-radius:10px;padding:14px;text-align:center">
            <div style="font-size:22px;font-weight:700;color:#155724">12 Credits</div>
            <div style="color:#155724;font-size:13px">45-minute session</div></div>""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ⚙️ Session Configuration")

    col1, col2 = st.columns([1, 1])

    with col1:
        duration = st.selectbox(
            "⏱️ Interview Duration",
            options=["15 minutes", "30 minutes", "45 minutes"],
            index=1,
            help="Longer sessions include more questions across all difficulty levels.",
            key="interview_duration_select"
        )
        st.session_state.interview_duration = duration

        jd = st.text_area(
            "📋 Job Description",
            height=200,
            placeholder="Paste the full job description here...",
            key="interview_jd_input",
            value=st.session_state.interview_jd,
        )
        st.session_state.interview_jd = jd

    with col2:
        uploaded = st.file_uploader(
            "📄 Upload Your Resume (PDF / DOCX)",
            type=["pdf", "docx"],
            key="interview_resume_upload"
        )

        st.markdown("#### 📊 What You'll Get")
        st.markdown("""
        <ul style="list-style:none;padding:0;margin:0">
          <li>✅ 5 General interview questions</li>
          <li>✅ 5 Technical questions (based on JD)</li>
          <li>✅ 5 Resume-based questions (from your CV)</li>
          <li>✅ <b>15 questions total</b> per session</li>
          <li>✅ AI Interviewer — question by question</li>
          <li>✅ Type or speak your answers</li>
          <li>✅ Per-question timer</li>
          <li>✅ AI evaluation on meaning, keywords, structure</li>
          <li>✅ Full feedback report (PDF + DOCX)</li>
          <li>✅ Suggested improved answers</li>
        </ul>
        """, unsafe_allow_html=True)

    # ─── Generate & Download Q&A Bank ────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 Step 1: Generate & Download Q&A Bank")
    st.markdown("*You can download the full Q&A bank (PDF/DOCX) before or after practicing.*")

    gen_col1, gen_col2 = st.columns([2, 1])

    if jd.strip() and uploaded:
        with gen_col1:
            if st.button("📚 Generate Q&A Bank + Start Practice Session",
                         type="primary", key="start_practice_btn",
                         use_container_width=True):
                credits_needed = DURATION_CREDITS[duration]
                if not check_access_fn(required_credits=credits_needed):
                    st.error(f"⚠️ You need {credits_needed} credits for a {duration} session. Please top up.")
                    return

                with st.spinner("🤖 AI is generating your personalized interview questions..."):
                    try:
                        resume_text = extract_resume_fn(uploaded)
                        st.session_state.interview_resume_text = resume_text

                        qa_bank = generate_structured_interview_qa(resume_text, jd, duration)
                        st.session_state.interview_qa_bank = qa_bank

                        flat = flatten_questions(qa_bank)
                        st.session_state.interview_questions_flat = flat
                        st.session_state.interview_current_idx = 0
                        st.session_state.interview_session_results = []
                        st.session_state.interview_report = None

                        # Deduct credits
                        user_email = st.session_state.user_data["email"]
                        deduct_credits_fn(user_email, credits_needed, feature="Interview Practice")

                        st.session_state.interview_phase = "session"
                        st.success(f"✅ {len(flat)} questions generated! {credits_needed} credits used. Starting session...")
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ Failed to generate questions: {str(e)}")

    else:
        st.info("👆 Please upload your resume and enter a job description to begin.")

    # Download-only option (if bank already exists)
    if st.session_state.interview_qa_bank:
        st.markdown("---")
        st.markdown("### 📁 Download Previously Generated Q&A Bank")
        bank = st.session_state.interview_qa_bank
        resume_text = st.session_state.interview_resume_text
        jd_text = st.session_state.interview_jd

        # Convert structured bank to flat text for export
        flat_text = _qa_bank_to_text(bank)
        try:
            pdf_buf, docx_buf = export_qa_fn(flat_text)
            dc1, dc2 = st.columns(2)
            with dc1:
                st.download_button("📥 Download Q&A PDF", data=pdf_buf,
                                   file_name="interview_qa_bank.pdf", mime="application/pdf",
                                   key="dl_qa_pdf_setup")
            with dc2:
                st.download_button("📥 Download Q&A DOCX", data=docx_buf,
                                   file_name="interview_qa_bank.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key="dl_qa_docx_setup")
        except Exception as e:
            st.warning(f"Export failed: {e}")


def _render_voice_recorder(session_key: str):
    """Render voice recorder with inline onclick handlers (main page = mic works).
    Stores transcript in query params → auto-fills text area on next rerun."""
    qp = st.query_params
    voice_val = qp.get("v", "")
    if voice_val and voice_val.strip():
        st.session_state[session_key] = voice_val
        try:
            del qp["v"]
        except Exception:
            pass
        st.rerun()

    uid = hash(session_key) % 100000
    # All JS is inside onclick attributes (script tags don't execute in markdown)
    st.markdown(f"""
    <style>
        .vr-{uid} .vb {{ padding:10px 24px; border-radius:24px; border:none; cursor:pointer; font-size:15px; font-weight:600; margin:2px; }}
        .vr-{uid} .vbr {{ background:linear-gradient(135deg,#e94560,#c0392b); color:#fff; }}
        .vr-{uid} .vbs {{ background:linear-gradient(135deg,#2ecc71,#16a085); color:#fff; }}
        .vr-{uid} .vst {{ font-size:13px; color:#555; margin-top:6px; min-height:20px; }}
        .vr-{uid} .h {{ display:none !important; }}
    </style>
    <div class="vr-{uid}" style="margin-bottom:8px">
        <button class="vb vbr" id="s-{uid}"
            onclick="
                var S=window.SpeechRecognition||window.webkitSpeechRecognition;
                if(!S){{document.getElementById('st-{uid}').innerHTML='Not supported';return;}}
                window.__r{uid}=new S(); window.__t{uid}='';
                var r=window.__r{uid}; r.continuous=true; r.interimResults=true; r.lang='en-US';
                r.onstart=function(){{
                    document.getElementById('s-{uid}').classList.add('h');
                    document.getElementById('p-{uid}').classList.remove('h');
                    document.getElementById('st-{uid}').innerHTML='🔴 Recording...';
                }};
                r.onresult=function(e){{for(var i=e.resultIndex;i<e.results.length;i++)if(e.results[i].isFinal)window.__t{uid}+=e.results[i][0].transcript+' ';}};
                r.onend=function(){{
                    document.getElementById('s-{uid}').classList.remove('h');
                    document.getElementById('p-{uid}').classList.add('h');
                    var t=(window.__t{uid}||'').trim();
                    if(t){{document.getElementById('st-{uid}').innerHTML='✅ Done!'; window.location.search='?v='+encodeURIComponent(t);}}
                    else document.getElementById('st-{uid}').innerHTML='No speech. Try again.';
                }};
                r.onerror=function(e){{
                    document.getElementById('s-{uid}').classList.remove('h');
                    document.getElementById('p-{uid}').classList.add('h');
                    document.getElementById('st-{uid}').innerHTML=e.error==='not-allowed'?'❌ Mic blocked':e.error==='no-speech'?'No speech.':'Error: '+e.error;
                }};
                try{{r.start();}}catch(e){{document.getElementById('st-{uid}').innerHTML='Error: '+e.message;}}
            ">🎙️ Start Speaking</button>
        <button class="vb vbs h" id="p-{uid}"
            onclick="if(window.__r{uid})window.__r{uid}.stop();">⏹️ Stop</button>
        <div class="vst" id="st-{uid}">Click "Start Speaking" to begin</div>
    </div>
    """, unsafe_allow_html=True)


def _timer_component_html(seconds: int, question_idx: int = 0) -> str:
    """HTML/JS countdown timer that shows remaining time per question. Resets per question via unique idx."""
    seed = question_idx % 10000
    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Courier New', monospace;
            display: flex; align-items: center; justify-content: center;
            min-height: 70px; background: transparent;
        }}
        .timer {{
            padding: 12px 24px; border-radius: 12px;
            background: linear-gradient(135deg, #1a1a2e, #16213e);
            color: #fff; font-size: 32px; font-weight: 700;
            text-align: center; letter-spacing: 2px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.2);
            min-width: 200px;
        }}
        .timer.warning {{ color: #e67e22; }}
        .timer.danger {{ color: #e94560; animation: blink 0.8s infinite; }}
        .label {{ font-size: 11px; color: #a0b4d6; text-transform: uppercase; letter-spacing: 1px; margin-top: 4px; }}
        @keyframes blink {{ 0%,100%{{opacity:1}} 50%{{opacity:0.3}} }}
    </style>
</head>
<body>
    <div>
        <div class="timer" id="timer_{seed}">{seconds // 60}:{seconds % 60:02d}</div>
        <div class="label">Q{question_idx+1} — Remaining</div>
    </div>
    <script>
    (function() {{
        var total = {seconds};
        var el = document.getElementById('timer_{seed}');
        function tick() {{
            if (total <= 0) return;
            total--;
            var m = Math.floor(total / 60);
            var s = total % 60;
            el.textContent = m + ':' + (s < 10 ? '0' : '') + s;
            el.className = 'timer' + (total <= 10 ? ' danger' : total <= 30 ? ' warning' : '');
            setTimeout(tick, 1000);
        }}
        tick();
    }})();
    </script>
</body>
</html>
    """


def _qa_bank_to_text(bank: dict) -> str:
    """Convert structured Q&A bank dict to displayable text for existing export_interview_qa."""
    lines = []
    idx = 1
    section_titles = {
        "general": "General Questions",
        "technical": "Technical Questions (JD-based)",
        "resume": "Resume-based Questions",
    }
    for section in ["general", "technical", "resume"]:
        section_title = section_titles.get(section, section.title())
        lines.append(f"\n=== {section_title} ===\n")
        qs = bank.get(section, [])
        for q in qs:
            lines.append(f"{idx}. {q.get('question', '')}")
            lines.append(f"Answer: {q.get('ideal_answer', '')}")
            lines.append("")
            idx += 1
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# Phase 2: Interview Session
# ─────────────────────────────────────────────────────────────────────────────
def _phase_session():
    questions = st.session_state.interview_questions_flat
    idx = st.session_state.interview_current_idx
    total = len(questions)

    if not questions:
        st.error("No questions found. Please go back and regenerate.")
        if st.button("↩️ Back to Setup"):
            st.session_state.interview_phase = "setup"
            st.rerun()
        return

    # Progress header
    progress_pct = idx / total
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:14px;padding:20px;margin-bottom:18px">
        <div style="display:flex;justify-content:space-between;align-items:center">
            <div>
                <h2 style="color:#e94560;margin:0">🤖 AI Interviewer</h2>
                <p style="color:#a0b4d6;margin:4px 0 0">Answer each question as if in a real interview</p>
            </div>
            <div style="text-align:right">
                <div style="color:#fff;font-size:20px;font-weight:700">Question {idx+1}/{total}</div>
                <div style="color:#a0b4d6;font-size:13px">Session in progress</div>
            </div>
        </div>
        <div style="margin-top:14px;background:#0f3460;border-radius:8px;height:8px;overflow:hidden">
            <div style="background:linear-gradient(90deg,#e94560,#fc5c7d);height:8px;width:{progress_pct*100:.1f}%;
                        border-radius:8px;transition:width 0.5s"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if idx >= total:
        # All questions answered
        _wrap_up_session()
        return

    q_obj = questions[idx]
    section = q_obj["section"].title()
    difficulty = q_obj["difficulty"]
    question_text = q_obj["question"]

    # Difficulty badge color
    diff_colors = {"Simple": "#2ecc71", "Hard": "#e67e22", "Very Hard": "#e94560"}
    diff_color = diff_colors.get(difficulty, "#888")

    # Determine per-question time limit (total duration in sec / number of questions)
    total_duration_map = {"15 minutes": 900, "30 minutes": 1800, "45 minutes": 2700}
    total_sec = total_duration_map.get(st.session_state.interview_duration, 1800)
    sec_per_q = max(60, total_sec // max(len(questions), 1))

    # Timer card
    st.markdown(f"""
    <div style="border-left:4px solid {diff_color};background:#f8f9ff;border-radius:0 12px 12px 0;
                padding:20px;margin-bottom:16px">
        <div style="display:flex;gap:10px;margin-bottom:10px;flex-wrap:wrap">
            <span style="background:#16213e;color:#a0b4d6;padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">
                {section}
            </span>
            <span style="background:#6c5ce7;color:#fff;padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600">
                🔊 Question Audio
            </span>
        </div>
        <p style="font-size:18px;font-weight:600;color:#1a1a2e;margin:0;line-height:1.5">{question_text}</p>
    </div>
    """, unsafe_allow_html=True)

    # Timer + TTS row
    timer_col, tts_col = st.columns([1, 1])
    with timer_col:
        st.components.v1.html(_timer_component_html(sec_per_q, idx), height=80)
    with tts_col:
        st.components.v1.html(tts_component_html(question_text), height=80)

    # ── Answer input ──────────────────────────────────────────────────────────
    st.markdown("#### ✍️ Your Answer")

    voice_key = f"voice_answer_{idx}"
    typed_key = f"typed_answer_{idx}"

    tab_type, tab_voice = st.tabs(["⌨️ Type Answer", "🎙️ Speak Answer"])

    with tab_type:
        typed_answer = st.text_area(
            "Type your answer here:",
            height=200,
            placeholder="Give a structured answer. Use STAR format. Be specific and detailed.",
            key=typed_key,
            label_visibility="collapsed"
        )

    with tab_voice:
        _render_voice_recorder(voice_key)

        voice_answer = st.text_area(
            "Your answer (editable):",
            height=140,
            placeholder="Transcript auto-fills here after speaking. Edit if needed...",
            key=voice_key,
            label_visibility="collapsed",
        )

    # Combine: typed preferred, fallback to voice
    final_answer = typed_answer.strip() if typed_answer.strip() else voice_answer.strip()

    # ── Skip / Submit ─────────────────────────────────────────────────────────
    btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])

    with btn_col1:
        submit_disabled = not final_answer
        if st.button("✅ Submit Answer & Next Question",
                     type="primary", key=f"submit_{idx}",
                     disabled=submit_disabled,
                     use_container_width=True):
            _submit_answer(q_obj, final_answer)

    with btn_col2:
        if st.button("⏭️ Skip Question", key=f"skip_{idx}", use_container_width=True):
            _submit_answer(q_obj, "")  # empty = skipped

    with btn_col3:
        if st.button("🔚 End Session Early", key=f"end_{idx}", use_container_width=True):
            _wrap_up_session(early=True)

    # ── Hints panel ───────────────────────────────────────────────────────────
    with st.expander("💡 Hints — Key Points to Cover", expanded=False):
        st.markdown("*These are the key concepts your answer should address (shown for practice purposes):*")
        for kp in q_obj.get("key_points", []):
            st.markdown(f"• {kp}")

    # ── Previous results quick view ───────────────────────────────────────────
    if st.session_state.interview_session_results:
        with st.expander(f"📋 Previous Answers — {len(st.session_state.interview_session_results)} answered", expanded=False):
            for i, res in enumerate(st.session_state.interview_session_results):
                ev = res.get("evaluation", {})
                score = ev.get("score", 0)
                color = "#2ecc71" if score >= 70 else "#e67e22" if score >= 50 else "#e94560"
                q_text = res["question_obj"]["question"][:80]
                st.markdown(f"""
                <div style="display:flex;justify-content:space-between;padding:8px;
                            border-radius:8px;background:#f8f9ff;margin-bottom:6px">
                    <span style="font-size:13px;color:#333">Q{i+1}. {q_text}...</span>
                    <span style="font-size:13px;font-weight:700;color:{color}">{score}/100</span>
                </div>
                """, unsafe_allow_html=True)


def _submit_answer(q_obj: dict, answer: str):
    """Evaluate answer and advance to next question."""
    with st.spinner("🤖 AI is evaluating your answer..."):
        evaluation = evaluate_answer(
            question=q_obj["question"],
            ideal_answer=q_obj["ideal_answer"],
            key_points=q_obj["key_points"],
            user_answer=answer,
            section=q_obj["section"],
            difficulty=q_obj["difficulty"],
        )

    st.session_state.interview_session_results.append({
        "question_obj": q_obj,
        "user_answer": answer,
        "evaluation": evaluation,
    })

    # Show quick score feedback
    score = evaluation.get("score", 0)
    if score >= 75:
        st.success(f"🎉 Great answer! Score: {score}/100")
    elif score >= 50:
        st.warning(f"👍 Good attempt. Score: {score}/100")
    else:
        st.error(f"📚 Needs improvement. Score: {score}/100")

    # Advance
    st.session_state.interview_current_idx += 1
    st.session_state.voice_transcript_buffer = ""

    total = len(st.session_state.interview_questions_flat)
    if st.session_state.interview_current_idx >= total:
        _wrap_up_session()
    else:
        st.rerun()


def _wrap_up_session(early: bool = False):
    """Generate report and move to report phase."""
    results = st.session_state.interview_session_results
    if not results:
        st.warning("No answers recorded. Please answer at least one question.")
        return

    if early:
        st.info(f"Session ended early. {len(results)} questions answered.")

    with st.spinner("📊 Generating your feedback report..."):
        report = generate_feedback_report(results, st.session_state.interview_duration)
        st.session_state.interview_report = report
        st.session_state.interview_phase = "report"

    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 3: Feedback Report
# ─────────────────────────────────────────────────────────────────────────────
def _phase_report():
    report = st.session_state.interview_report
    if not report:
        st.error("No report available.")
        return

    score = report["overall_score"]
    band = report["performance_band"]
    band_color_map = {"Excellent": "#2ecc71", "Good": "#3498db", "Average": "#e67e22", "Needs Improvement": "#e94560"}
    band_color = band_color_map.get(band, "#888")

    # ── Score card ────────────────────────────────────────────────────────────
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1a1a2e,#16213e);border-radius:16px;padding:28px;margin-bottom:20px;text-align:center">
        <h1 style="color:#e94560;margin:0;font-size:52px;font-weight:800">{score}<span style="font-size:24px">/100</span></h1>
        <div style="color:{band_color};font-size:22px;font-weight:700;margin:6px 0">{band}</div>
        <p style="color:#a0b4d6;margin:0">Your Interview Performance Score</p>
    </div>
    """, unsafe_allow_html=True)

    # Section scores
    mc1, mc2, mc3, mc4, mc5 = st.columns(5)
    with mc1:
        st.metric("📋 General", f"{report['general_score']}/100")
    with mc2:
        st.metric("⚙️ Technical", f"{report['technical_score']}/100")
    with mc3:
        st.metric("📄 Resume", f"{report['resume_score']}/100")
    with mc4:
        st.metric("✅ Questions Done", report["total_questions"])
    with mc5:
        st.metric("✅ Answered Well", len(report["well_answered"]))

    st.markdown("---")

    # ── Overall Summary ───────────────────────────────────────────────────────
    st.markdown("### 📋 Overall Assessment")
    st.info(report.get("overall_summary", ""))

    col_s, col_w = st.columns(2)
    with col_s:
        st.markdown("#### 💪 Strengths")
        for s in report.get("key_strengths", []):
            st.markdown(f"✅ {s}")

    with col_w:
        st.markdown("#### ⚠️ Areas to Improve")
        for w in report.get("weak_areas", []):
            st.markdown(f"🔸 {w}")

    # ── Section feedback ──────────────────────────────────────────────────────
    col_g, col_t, col_r = st.columns(3)
    with col_g:
        st.markdown("#### 📋 General")
        st.markdown(report.get("general_feedback", ""))
    with col_t:
        st.markdown("#### ⚙️ Technical")
        st.markdown(report.get("technical_feedback", ""))
    with col_r:
        st.markdown("#### 📄 Resume")
        st.markdown(report.get("resume_feedback", ""))

    st.markdown("---")

    # ── Per-question results ──────────────────────────────────────────────────
    st.markdown("### 📊 Question-by-Question Results")

    for i, res in enumerate(report.get("session_results", [])):
        q_obj = res.get("question_obj", {})
        ev = res.get("evaluation", {})
        user_ans = res.get("user_answer", "")
        s = ev.get("score", 0)
        card_color = "#d4edda" if s >= 70 else "#fff3cd" if s >= 50 else "#f8d7da"
        border_color = "#28a745" if s >= 70 else "#ffc107" if s >= 50 else "#dc3545"

        with st.expander(f"Q{i+1}. {q_obj.get('question','')[:80]}... — {s}/100", expanded=False):
            st.markdown(f"""
            <div style="border-left:4px solid {border_color};background:{card_color};
                        border-radius:0 10px 10px 0;padding:14px;margin-bottom:12px">
                <b>{q_obj.get('section','').title()} | {q_obj.get('difficulty','')}</b><br>
                {q_obj.get('question','')}
            </div>
            """, unsafe_allow_html=True)

            tc1, tc2, tc3, tc4 = st.columns(4)
            tc1.metric("Overall", f"{ev.get('score',0)}/100")
            tc2.metric("Meaning Match", f"{ev.get('meaning_match',0)}/100")
            tc3.metric("Keyword Coverage", f"{ev.get('keyword_coverage',0)}/100")
            tc4.metric("Structure", f"{ev.get('structure_score',0)}/100")

            if user_ans:
                st.markdown("**Your Answer:**")
                st.info(user_ans)
            else:
                st.warning("*Question skipped*")

            kc1, kc2 = st.columns(2)
            with kc1:
                st.markdown("**✅ Keywords Covered:**")
                covered = ev.get("keywords_covered", [])
                st.markdown(", ".join(covered) if covered else "_None_")
            with kc2:
                st.markdown("**❌ Keywords Missed:**")
                missed = ev.get("keywords_missed", [])
                st.markdown(", ".join(missed) if missed else "_None — great!_")

            st.markdown("**💡 Coach Feedback:**")
            st.markdown(ev.get("brief_feedback", ""))

            if ev.get("strengths"):
                st.markdown("**Strengths:**")
                for s_item in ev["strengths"]:
                    st.markdown(f"✔ {s_item}")

            if ev.get("improvements"):
                st.markdown("**Improvements:**")
                for imp in ev["improvements"]:
                    st.markdown(f"• {imp}")

            if ev.get("improved_answer"):
                st.markdown("**📝 Suggested Improved Answer:**")
                st.success(ev["improved_answer"])

    st.markdown("---")

    # ── Keywords summary ──────────────────────────────────────────────────────
    st.markdown("### 🔑 Keyword Summary")
    kw_col1, kw_col2 = st.columns(2)
    with kw_col1:
        st.markdown("**Keywords Covered:**")
        covered_kws = report.get("keywords_covered", [])
        if covered_kws:
            badges = " ".join([f'<span style="background:#d4edda;color:#155724;padding:3px 10px;border-radius:15px;font-size:12px;margin:2px;display:inline-block">{k}</span>' for k in covered_kws])
            st.markdown(badges, unsafe_allow_html=True)
        else:
            st.markdown("_None_")
    with kw_col2:
        st.markdown("**Keywords Missed:**")
        missed_kws = report.get("keywords_missed", [])
        if missed_kws:
            badges = " ".join([f'<span style="background:#f8d7da;color:#721c24;padding:3px 10px;border-radius:15px;font-size:12px;margin:2px;display:inline-block">{k}</span>' for k in missed_kws])
            st.markdown(badges, unsafe_allow_html=True)
        else:
            st.markdown("_None — excellent coverage!_ 🎉")

    st.markdown("---")

    # ── Recommendations ───────────────────────────────────────────────────────
    st.markdown("### 🎯 Recommendations For Further Practice")
    for rec in report.get("recommendations", []):
        st.markdown(f"▶ {rec}")

    st.markdown("### 🚀 Next Steps")
    st.info(report.get("next_steps", ""))

    st.markdown("---")

    # ── Downloads ─────────────────────────────────────────────────────────────
    st.markdown("### 📥 Download Your Feedback Report")
    try:
        with st.spinner("Preparing report files..."):
            pdf_buf, docx_buf = export_feedback_report(report)

        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("📄 Download PDF Report", data=pdf_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                               mime="application/pdf", key="dl_report_pdf", use_container_width=True)
        with dl2:
            st.download_button("📝 Download DOCX Report", data=docx_buf,
                               file_name=f"interview_feedback_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="dl_report_docx", use_container_width=True)
    except Exception as e:
        st.error(f"Report export failed: {e}")

    st.markdown("---")

    # ── Restart ───────────────────────────────────────────────────────────────
    rc1, rc2 = st.columns(2)
    with rc1:
        if st.button("🔁 Practice Again (New Session)", type="primary", use_container_width=True):
            _reset_interview_session()
            st.rerun()
    with rc2:
        if st.button("⚙️ Change Settings", use_container_width=True):
            st.session_state.interview_phase = "setup"
            st.rerun()


def _reset_interview_session():
    keys = [
        "interview_phase", "interview_qa_bank", "interview_questions_flat",
        "interview_current_idx", "interview_session_results", "interview_report",
        "voice_transcript_buffer",
    ]
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]
